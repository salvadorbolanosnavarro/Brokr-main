from fastapi import (FastAPI, HTTPException, Query, Request, UploadFile, File,
                     BackgroundTasks, Response)
from fastapi.middleware.cors import CORSMiddleware
from limites import exigir_cupo, exigir_sesion
from pydantic import BaseModel
from core.auth import get_user_id_from_token
from core.config import settings
from core.database import call_public_rpc, call_service_rpc, delete_rows, get_public_rows, get_rows, get_service_json, get_service_json_or_empty, patch_rows, patch_rows_ignoring_http_status, patch_rows_no_response, post_rows, upsert_rows
from core.legacy_main_config import legacy_main_settings
from core.legacy_admin import require_legacy_admin as require_admin
from core.telemetry import (_request_modulo, _track_anthropic, _track_gemini_image, _track_groq, track_usage)
from core.user_access import get_user_access_state, get_user_rol
from core.subscriptions import (expire_trial_subscription as _expirar_trial_suscripcion, trial_has_expired as _trial_ya_vencio, trial_max_available as _trial_max_disponible)
from core.stripe import (
    EMPRESA_ASIENTOS_BASE, EMPRESA_ASIENTOS_MAX, EMPRESA_TARIFAS,
    PROMO_CODE_AMPI, STRIPE_PRICE_AMPI, STRIPE_PRICE_PRO,
    STRIPE_PRICE_EMPRESA_ANUAL, STRIPE_PRICE_EMPRESA_EXTRA_ANUAL,
    STRIPE_PRICE_EMPRESA_EXTRA_MENSUAL, STRIPE_PRICE_EMPRESA_MENSUAL,
    STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET, TRIAL_MAX_DIAS,
    get_or_create_stripe_customer as _get_or_create_stripe_customer,
    precio_empresa as _precio_empresa, stripe_headers as _stripe_headers,
)
from core.facebook_tokens import facebook_token_state as _fb_estado_token
from core.cache import cache_get, cache_set
from core.contact_import import map_org_agents as _mapa_agentes_org
from core.easybroker import EB_API_KEY, EB_BASE, _EB_LOTE, _EB_PAUSA_LOTE, _eb_get_reintentos, eb_headers, extract_colonia, normalize
from core.easybroker_mapping import _EB_LIMITE_PROPIEDADES, _EB_STATUS_DEFAULT, _EB_STATUS_MAP, _eb_to_brokr
from core.easybroker_migration import set_import_progress as _prog
from core.pdf_design import theme_css_for_pdf
from core.pdf_store import _pdf_store
import httpx
import os
import time
import re
import asyncio
import logging
import base64
import hmac
import hashlib
import uuid as _uuid
import io
import json
from typing import Optional, List, Dict, Any
from datetime import datetime, date, timedelta, timezone
from pathlib import Path

try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

# Pillow


from routers.admin_usage import router as admin_usage_router
from routers.account_delete import router as account_delete_router
from routers.avm_legacy import router as avm_legacy_router

from routers.avm_claude import router as avm_claude_router

from routers.avm_websearch import router as avm_websearch_router

from core.facebook_tokens import FACEBOOK_REQUIRED_SCOPES
from routers.facebook_connection_read import router as facebook_connection_read_router

from core.facebook_secrets import (decrypt_facebook_secret as descifrar_secreto, encrypt_facebook_secret as cifrar_secreto, facebook_secret_encryption_available)

from core.facebook_connection_store import get_facebook_meta_row as _fb_get_meta_row

from core.facebook_graph import (
    FB_API_VERSION,
    FB_GRAPH,
    _FB_CODIGOS_REINTENTABLES,
    _FB_CODIGOS_TOKEN,
    _FB_ERRORES_COMUNES,
    _FB_ESPERA_BASE,
    _FB_ESPERA_MAX,
    _FB_REINTENTOS,
    _FB_USAR_PROOF,
    _fb_appsecret_proof,
    _fb_debe_reintentar,
    _fb_espera_por_uso,
    _fb_exigir_ok,
    _fb_friendly_error,
    _fb_get_json,
    _fb_parse_error,
    _fb_batch,
)

from routers.facebook_pages import router as facebook_pages_router

from core.facebook_connection_store import patch_facebook_meta as _fb_patch_meta

from routers.facebook_select_page import router as facebook_select_page_router

from routers.facebook_select_ad_account import router as facebook_select_ad_account_router

from routers.facebook_encrypt_tokens import router as facebook_encrypt_tokens_router

from core.facebook_token_lifecycle import (FB_TOKEN_DEFAULT_LIFETIME_SECONDS as _FB_TOKEN_VIDA_DEFECTO, debug_facebook_token as _fb_debug_token)

from routers.facebook_refresh_token import router as facebook_refresh_token_router

from routers.facebook_disconnect import router as facebook_disconnect_router


from routers.facebook_ad_accounts import router as facebook_ad_accounts_router

from routers.facebook_city_search import router as facebook_city_search_router

from core.facebook_insights import (
    FB_BREAKDOWNS as _FB_BREAKDOWNS,
    FB_DATE_PRESETS as _FB_DATE_PRESETS,
    FB_INSIGHTS_FIELDS as _FB_INSIGHTS_FIELDS,
    FB_KEY_ACTIONS as _FB_ACCIONES_CLAVE,
    normalize_facebook_insights as _fb_normaliza_insights,
)


from routers.facebook_campaigns import router as facebook_campaigns_router

from routers.facebook_insights_read import router as facebook_insights_read_router

from routers.facebook_campaign_review import router as facebook_campaign_review_router

from core.facebook_leadgen_config import (
    FB_VERIFY_TOKEN,
    FB_WEBHOOK_SECRET as _FB_WEBHOOK_SECRET,
)

from routers.facebook_leadgen_verify import router as facebook_leadgen_verify_router

from routers.facebook_leadgen_status import router as facebook_leadgen_status_router

from routers.facebook_leadgen_subscribe import router as facebook_leadgen_subscribe_router

from core.facebook_persistence import (
    FACEBOOK_AD_ENTITIES_TABLE as _FB_TABLA_ENTIDADES,
    facebook_table_missing as _fb_tabla_falta,
    warn_facebook_migration as _fb_avisa_migracion,
)

from core.facebook_leadgen_processor import (
    FACEBOOK_LEAD_FIELDS as _FB_CAMPOS_LEAD,
    find_facebook_page_owner as _fb_buscar_dueno_de_pagina,
    process_facebook_lead as _fb_procesar_lead,
)

from routers.facebook_leadgen_webhook import router as facebook_leadgen_webhook_router

from routers.facebook_page_posts import router as facebook_page_posts_router

from routers.facebook_audiences_read import router as facebook_audiences_read_router

from routers.facebook_oauth_callback import router as facebook_oauth_callback_router

from routers.facebook_publish import router as facebook_publish_router

from routers.facebook_publish_property import router as facebook_publish_property_router

from routers.facebook_save_page import router as facebook_save_page_router

from routers.facebook_ad_description import router as facebook_ad_description_router

from routers.facebook_campaign_toggle import router as facebook_campaign_toggle_router

from core.facebook_persistence import (
    find_facebook_creation_by_idempotency as _fb_buscar_por_idempotencia,
    reserve_facebook_creation as _fb_reservar_creacion,
    update_facebook_entity as _fb_actualizar_entidad,
)

from routers.facebook_reconcile import router as facebook_reconcile_router

from routers.facebook_audiences import router as facebook_audiences_router

from routers.facebook_create_ad import router as facebook_create_ad_router
from routers.facebook_qa_selfcheck import router as facebook_qa_selfcheck_router
app = FastAPI()
app.include_router(facebook_qa_selfcheck_router)
app.include_router(facebook_create_ad_router)

app.include_router(facebook_audiences_router)

app.include_router(facebook_reconcile_router)

app.include_router(facebook_refresh_token_router)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# whatsapp.py es el módulo de WhatsApp (multi-número, IA de recepción, webhook
# propio bajo /whatsapp2 — el prefijo interno del router no cambió aunque el
# archivo ya se llama whatsapp.py). Import defensivo: si algo le falta, el
# resto del backend sigue vivo.
try:
    from whatsapp import router as whatsapp_router
    app.include_router(whatsapp_router)
except Exception as _e:
    import logging as _logging
    _logging.getLogger("broquer.main").error("No se pudo cargar whatsapp: %s", _e)

# Motor agéntico de Broq (tool-use nativo + loop de varios pasos + voz Whisper).
# Import defensivo: si por cualquier razón fallara la carga, el resto del backend
# sigue funcionando con normalidad.
try:
    from routers.agente import router as agente_router
    app.include_router(agente_router)
except Exception as _e:
    print(f"[agente] No se pudo montar el router agéntico: {_e}")

# Broquer para empresas: miembros, invitaciones, roles y permisos.
# Mismo import defensivo: si falla, el resto del backend sigue vivo.
try:
    from routers.organizaciones import router as org_router
    app.include_router(org_router)
except Exception as _e:
    print(f"[org] No se pudo montar el router de organizaciones: {_e}")

# WhatsApp de ChatGPT: onboarding real por Meta Embedded Signup, separado del módulo legacy.
try:
    from routers.whatsapp_chatgpt import router as whatsapp_chatgpt_router
    app.include_router(whatsapp_chatgpt_router)
except Exception as _e:
    print(f"[whatsapp-chatgpt] No se pudo montar el router: {_e}")
# Cumplimiento PLD/UIF: expediente único, umbrales, avisos y bitácora.
# Mismo import defensivo: si falla, el resto del backend sigue vivo.
try:
    from routers.cumplimiento import router as pld_router
    app.include_router(pld_router)
except Exception as _e:
    print(f"[pld] No se pudo montar el router de cumplimiento: {_e}")

# Firma electrónica: documentos, firmantes, código de verificación, constancia
# y verificación pública por folio. Mismo import defensivo: si falla, el resto
# del backend sigue vivo.
try:
    from routers.firmas import router as firmas_router
    app.include_router(firmas_router)
except Exception as _e:
    print(f"[firmas] No se pudo montar el router de firma electrónica: {_e}")

# Video de ficha: arma un recorrido con ffmpeg a partir de las fotos que ya
# viven en la propiedad. Mismo import defensivo: si falla, el resto del
# backend sigue vivo.
try:
    from routers.video import router as video_router
    app.include_router(video_router)
except Exception as _e:
    print(f"[video] No se pudo montar el router de video: {_e}")

# Correo electrónico: conexión IMAP/SMTP, bandeja, lectura, respuesta y
# envío. Mismo import defensivo: si falla, el resto del backend sigue vivo.
try:
    from routers.correo import router as correo_router
    app.include_router(correo_router)
except Exception as _e:
    print(f"[correo] No se pudo montar el router de correo: {_e}")

# Bolsa inmobiliaria: inventario compartido entre agentes Broquer con
# comisión compartida. Mismo import defensivo: si falla, el resto del
# backend sigue vivo.
try:
    from routers.bolsa import router as bolsa_router
    app.include_router(bolsa_router)
except Exception as _e:
    print(f"[bolsa] No se pudo montar el router de la bolsa: {_e}")

# Finanzas: cuentas, ingresos, gastos, rentabilidad por propiedad, lectura
# de tickets con Broq y reportes PDF/CSV. Mismo import defensivo: si falla,
# el resto del backend sigue vivo.
try:
    from routers.finanzas import router as finanzas_router
    app.include_router(finanzas_router)
except Exception as _e:
    print(f"[finanzas] No se pudo montar el router de finanzas: {_e}")

# Solicitud pública de demos.
from routers.demo import router as demo_router
app.include_router(demo_router)

# Cuadrícula pública de Instagram para el landing.
from routers.instagram import router as instagram_router
app.include_router(instagram_router)

# Captura pública de leads desde los sitios de agentes.
from routers.public_site_leads import router as public_site_leads_router
app.include_router(public_site_leads_router)

# Estado mínimo del servicio.
from routers.system import router as system_router
app.include_router(system_router)

# INPC y UDIS desde Banxico SIE.
from routers.banxico import router as banxico_router
app.include_router(banxico_router)

# Heartbeat de uso por módulo.
from routers.telemetry import router as telemetry_router
app.include_router(telemetry_router)

# Proxy de chat Groq.
from routers.chat import router as chat_router
app.include_router(chat_router)

# Configuración pública para el frontend.
from routers.public_config import router as public_config_router
app.include_router(public_config_router)

# Conexión EasyBroker compartida por organización.
from routers.easybroker_config import get_eb_key_for_user, router as easybroker_config_router
app.include_router(easybroker_config_router)

# Importación del historial de leads de EasyBroker.
from routers.easybroker_import_stats import router as easybroker_import_stats_router
app.include_router(easybroker_import_stats_router)

# Coordinador de migración completa EasyBroker.
from routers.easybroker_migration import router as easybroker_migration_router
app.include_router(easybroker_migration_router)

# Importación de contactos directamente desde EasyBroker.
from routers.easybroker_contact_import import router as easybroker_contact_import_router
app.include_router(easybroker_contact_import_router)

# Estado de migración de fotos EasyBroker.
from routers.easybroker_photo_status import (_migrar_fotos_org, router as easybroker_photo_status_router)
app.include_router(easybroker_photo_status_router)

# Diagnóstico de la API de EasyBroker (solo lectura).
from routers.easybroker_diagnostics import router as easybroker_diagnostics_router
app.include_router(easybroker_diagnostics_router)

# Lectura autenticada de propiedades EasyBroker.
from routers.easybroker_properties import router as easybroker_properties_router
app.include_router(easybroker_properties_router)

# Listado legacy de propiedades EasyBroker (solo lectura).
from routers.easybroker_catalog import router as easybroker_catalog_router
app.include_router(easybroker_catalog_router)

# Autocomplete de colonias desde EasyBroker.
from routers.easybroker_colonias import router as easybroker_colonias_router
app.include_router(easybroker_colonias_router)

# Descargas de PDFs generados en memoria.
from routers.pdf_downloads import router as pdf_downloads_router
app.include_router(pdf_downloads_router)

# Generación de PDF para ISR.
from routers.isr_pdf import router as isr_pdf_router
app.include_router(isr_pdf_router)

# Noticias inmobiliarias RSS.
from routers.news import router as news_router
app.include_router(news_router)

# Descripción IA para ficha manual.
from routers.ficha_manual import router as ficha_manual_router
app.include_router(ficha_manual_router)

# Comparables AVM vía Apify/Inmuebles24.
from routers.avm_apify import router as avm_apify_router
app.include_router(avm_apify_router)

# Colonias AVM vía Google Places.
from routers.avm_places import router as avm_places_router
app.include_router(avm_places_router)

# Comparables AVM cercanos vía Supabase/PostGIS.
from routers.avm_nearby import router as avm_nearby_router
app.include_router(avm_nearby_router)

# Limpieza y edición de imágenes inmobiliarias.
from routers.image_cleaner import router as image_cleaner_router
app.include_router(image_cleaner_router)

# Recordatorios de tareas/citas en background.
from routers.reminders import router as reminders_router
app.include_router(reminders_router)

# Generación de contrato DOCX estándar.
from routers.contracts_basic import router as contracts_basic_router
app.include_router(contracts_basic_router)

# Contratos personalizados (machotes).
from routers.machotes import router as machotes_router
app.include_router(machotes_router)

# Compatibility aliases while main.py is progressively decomposed. All runtime
# environment names and public/privileged Supabase key policy live in Core.
GROQ_API_KEY     = settings.groq_api_key
ANTHROPIC_API_KEY = settings.anthropic_api_key
GEMINI_API_KEY    = settings.gemini_api_key
GROQ_BASE        = "https://api.groq.com/openai/v1"
ANTHROPIC_BASE   = "https://api.anthropic.com/v1"
GEMINI_BASE      = "https://generativelanguage.googleapis.com/v1beta"
APIFY_API_KEY = settings.apify_api_key
GOOGLE_PLACES_KEY = settings.google_places_key
SUPABASE_URL      = settings.supabase_url
SUPABASE_KEY      = settings.supabase_anon_key
FB_APP_ID     = settings.legacy_main_fb_app_id
FB_APP_SECRET = settings.legacy_main_fb_app_secret
FRONTEND_URL  = settings.legacy_main_frontend_url
# Banxico SIE — INPC + UDIS para calculadora ISR
BANXICO_TOKEN     = settings.banxico_token
BANXICO_BASE      = "https://www.banxico.org.mx/SieAPIRest/service/v1/series"
BANXICO_SERIE_UDIS = settings.banxico_series_udis  # Valor de UDIS (diaria)
BANXICO_SERIE_INPC = settings.banxico_series_inpc  # INPC mensual base 2Q-jul-2018=100
# service_role key — bypasea RLS. Solo para operaciones del backend en nombre
# del usuario, DESPUÉS de validar su JWT con get_user_id_from_token().
# NUNCA expongas esta variable al frontend.
SUPABASE_SERVICE_KEY = settings.supabase_service_key
# Pagos — Stripe

# ════════════════════════════════════════════════════════════════
# CONTEXTO DE ORGANIZACIÓN (Broquer para empresas)
# Tras la migración, la RLS filtra por org_id — NO por user_id. Todo registro
# que cree el backend debe llevar org_id o queda huérfano e invisible para
# todos. El backend usa service key y se brinca la RLS, así que un olvido aquí
# no truena: silenciosamente crea basura. Por eso va explícito en cada INSERT.
# ════════════════════════════════════════════════════════════════
from routers.organizaciones import (
    get_org_id_for_user, get_org_context, permiso_efectivo,
    exigir_gestion_integraciones,
)


# Suscripción de Broquer para Empresas.
from routers.subscription_enterprise import router as subscription_enterprise_router
app.include_router(subscription_enterprise_router)

# Checkout web de suscripción individual.
from routers.subscription_checkout import router as subscription_checkout_router
app.include_router(subscription_checkout_router)

# Cancelación de suscripción web.
from routers.subscription_cancel import router as subscription_cancel_router
app.include_router(subscription_cancel_router)

# Activación interna de suscripciones.
from routers.subscription_activate import router as subscription_activate_router
app.include_router(subscription_activate_router)

# Webhook de suscripciones web vía Stripe.
from routers.stripe_webhook import router as stripe_webhook_router
app.include_router(stripe_webhook_router)

# Webhook de suscripciones iOS vía RevenueCat.
from routers.revenuecat import router as revenuecat_router
app.include_router(revenuecat_router)

# Estado de suscripción y trial de Broquer Max.
from routers.subscription_status import router as subscription_status_router
app.include_router(subscription_status_router)

# Eliminación administrativa total (aislada; nunca se invoca en la auditoría).
from routers.admin_delete import router as admin_delete_router
app.include_router(admin_delete_router)

# Mutaciones administrativas no destructivas.
from routers.admin_accounts import router as admin_accounts_router
app.include_router(admin_accounts_router)

# Lecturas administrativas legacy.
from routers.admin_read import router as admin_read_router
app.include_router(admin_read_router)

# Estado unificado de perfil e integraciones.
from routers.profile_status import router as profile_status_router
app.include_router(profile_status_router)

# ────────────────────────────────────────────
# CLAUDE CHAT PROXY — BROQ IA SUPERINTELIGENTE
# ────────────────────────────────────────────
SHAARK_SYSTEM_PROMPT = """Eres Broq, el asistente de inteligencia artificial de la plataforma Broquer — el copiloto operativo para agentes inmobiliarios de México, especializada en Morelia y Michoacán.

IDENTIDAD:
- Tu nombre es Broq. Si el usuario dice "broq", "broker", "Broker", "broquer" o variantes, siempre escríbelo como "Broq" en tu respuesta.
- Eres el copiloto del agente. Puedes hacer casi todo lo que el agente haría manualmente en la plataforma — y lo haces por él cuando te lo pide.
- Eres especialmente útil cuando el agente va manejando, está en una cita, o no puede escribir. Si habla por voz, respondes con oraciones cortas y directas.
- Llamas al usuario por su nombre de pila cuando lo conoces (lo recibes en el contexto).

PERSONALIDAD:
- Hablas español mexicano, natural, cercano y profesional.
- Eres directo y preciso. Sin relleno. Sin redundancia.
- Nunca inventas cifras, leyes, artículos o datos que no existen.
- Si no sabes algo con certeza, lo dices y ofreces buscar o recomendar dónde verificar.

CONOCIMIENTO EXPERTO QUE DOMINAS:

DERECHO INMOBILIARIO MEXICANO:
- Código Civil Federal y de Michoacán: compraventa, arrendamiento, promesa de venta, comodato, cesión de derechos.
- Cuándo se requiere escritura pública ante notario y cuándo basta un contrato privado.
- Registro Público de la Propiedad: cómo registrar, por qué importa, tiempos y costos.
- Ley Federal de Protección de Datos Personales (LFPDPPP) — obligaciones del agente.
- Ley Federal para la Prevención e Identificación de Operaciones con Recursos de Procedencia Ilícita (LFPIORPI) — PLD para agentes inmobiliarios: reportes, aviso SAT, umbrales.
- Diferencias entre promesa de compraventa y contrato de compraventa definitivo.
- Derechos y obligaciones de arrendador y arrendatario: depósito, fianza, rescisión.
- Régimen de propiedad en condominio en Michoacán.
- Fideicomiso inmobiliario básico.
- Reglamentos de construcción de Morelia.

FISCAL E ISR:
- LISR artículos 119 y 120 — enajenación de inmuebles, exención 700,000 UDIS para casa habitación.
- Deducciones: precio de compra actualizado con INPC, mejoras, escrituración, comisiones.
- Retención del notario, declaración anual del vendedor.
- Régimen de arrendamiento en SAT: pagos provisionales, deducción ciega del 35%.
- ISAI (Impuesto Sobre Adquisición de Inmuebles) — quién lo paga, cuánto, dónde.
- IVA en operaciones comerciales e industriales.

VALUACIÓN Y MERCADO:
- Método de mercado (comparables), método físico (costo), capitalización de rentas.
- Cap rate, precio por m², análisis hedónico.
- Mercado de Morelia: Chapultepec, Altozano, Félix Ireta, Lomas del Estadio, Santa María, Lomas de Tzompantle, Vistas del Campestre, Villas del Pedregal, Bosques de Tariacuri, Torremolinos, Las Américas, Jardines del Rincón, y más.
- Factores de plusvalía: vialidades, equipamiento urbano, densidad, tendencia de zona.

MARKETING INMOBILIARIO:
- Facebook Ads e Instagram Ads para inmuebles: objetivos, presupuestos, públicos, creativos.
- Cómo redactar una ficha técnica que vende.
- Estrategia de precios: precio de lista vs precio de mercado.
- Cómo manejar la objeción de precio con el propietario.
- Técnicas de captación de exclusivas.
- Script de llamada en frío para propietarios.
- Presentación de servicios ante propietario.
- Marketing de contenidos: LinkedIn, Instagram, TikTok para agentes.

TECNOLOGÍA PARA AGENTES:
- EasyBroker: cómo conectar, importar propiedades, subir propiedades, el CRM.
- Portales: Inmuebles24, Vivanuncios, Lamudi, MercadoLibre Inmuebles.
- Firma electrónica en México: validez legal, Mifiel, Docusign.
- WhatsApp Business, Google Business Profile, Google Meet para agentes.
- Cómo usar Broquer al 100%: todos los módulos, cómo pedir ayuda por voz, etc.

CÓMO CONECTAR EASYBROKER (respuesta exacta cuando te pregunten):
1. En EasyBroker, haz clic en tu nombre (esquina superior derecha) → "Configuración de cuenta".
2. En el menú izquierdo, busca "Integraciones" o "API".
3. Copia tu API Key personal (código alfanumérico largo).
4. En Broquer, abre tu perfil haciendo clic en tus iniciales (esquina inferior izquierda del sidebar en desktop, o el avatar en móvil).
5. En la sección "EasyBroker", pega tu API Key y haz clic en "Conectar EasyBroker".
6. Broquer valida la conexión en segundos.
Nota: cada agente debe usar su propia API Key personal. No la compartas.

REGLA DE ORO PARA ACCIONES:
Cuando el usuario pide ejecutar una tarea, recopila los datos OBLIGATORIOS de UNO EN UNO, conversacionalmente. NUNCA ejecutes la acción con datos incompletos. Cuando tengas todo, di un resumen breve y ejecuta. Los opcionales que el usuario no conozca: usa 0 o "".

═══════════════════════════════════════════════════════════════
MODO ASISTENTE EJECUTOR — PRIORIDAD #1
═══════════════════════════════════════════════════════════════
Eres un ASISTENTE que EJECUTA, no un chatbot que sugiere. Cuando el usuario
pide algo que puedes hacer DIRECTAMENTE, HAZLO. No le digas "ve a tal módulo
y dale al botón X". TÚ lo haces y le entregas el resultado.

PREFIERE SIEMPRE LAS ACCIONES DIRECTAS sobre las que navegan:
  • `calcular_isr_directo`     → genera y descarga el PDF de ISR en el chat
  • `estimar_valor_directo`    → genera y descarga el PDF de estimación de valor
  • `agregar_contacto`         → agrega contacto al CRM sin salir del chat
  • `generar_contrato_directo` → descarga DOCX del contrato sin salir del chat

Solo navega (`llenar_isr`, `llenar_avm`, `llenar_contrato`, `navegar`) cuando:
  - El usuario explícitamente lo pide ("llévame a", "abre", "muéstrame el módulo de").
  - Faltan datos críticos y necesita editar a mano.

Tono: decidido, breve, fáctico. Di "Listo, lo hago." en lugar de "Voy a llevarte
a la pantalla de…". El usuario está manejando, dándote órdenes por voz; tú
ejecutas como una secretaria experta que conoce su trabajo.

══════════════════════════════════════════════════
ACCIÓN 1: CALCULAR ISR POR ENAJENACIÓN
══════════════════════════════════════════════════
Datos OBLIGATORIOS (pregunta uno por uno):
1. Tipo de inmueble: casa habitación, terreno, o comercial
2. Precio de venta (MXN)
3. Mes y año de la venta
4. Precio de compra original (MXN)
5. Mes y año de la compra
6. Si es casa: ¿usó la exención en los últimos 3 años? (sí / no / no sabe)
7. ¿Mejoras o ampliaciones? (monto o "no")
8. ¿Escrituración al comprar? (monto o "no sé")
9. ¿Comisión del agente en esta venta? (monto o "no aplica")

La pregunta 6 SOLO aplica a casa/departamento. Para terrenos y comerciales usa "no" automáticamente.

Cuando tengas todo:
[ACCION]{"tipo":"llenar_isr","precio_venta":NUMERO,"precio_compra":NUMERO,"anio_venta":NUMERO,"mes_venta":NUMERO,"anio_compra":NUMERO,"mes_compra":NUMERO,"inmueble":"casa","exencion":"no","mejoras":NUMERO,"escrituracion":NUMERO,"comision":NUMERO}[/ACCION]

Valores "inmueble": "casa" | "terreno" | "comercial"
Valores "exencion": "no" | "si" | "nose"
mes_venta y mes_compra son números 1-12. Datos opcionales desconocidos = 0.

══════════════════════════════════════════════════
ACCIÓN 2: OPINIÓN DE VALOR CON BÚSQUEDA WEB
══════════════════════════════════════════════════
Datos OBLIGATORIOS (pregunta uno por uno si faltan):
1. Colonia o fraccionamiento
2. Tipo de inmueble: casa, departamento, terreno, local, oficina, bodega
3. Operación: venta o renta
4. Superficie: m² construcción (casas/deptos/locales) o m² terreno (terrenos)
Opcionales: recámaras, baños, estacionamientos, condición terreno, ciudad (default Morelia).

[ACCION]{"tipo":"opinion_valor_web","colonia":"Vistas Altozano","tipo_inmueble":"terreno","operacion":"venta","m2_terreno":183,"m2_construccion":0,"recamaras":0,"banos":0,"ciudad":"Morelia","condicion_terreno":"plano"}[/ACCION]

Valores "tipo_inmueble": "casa" | "departamento" | "terreno" | "local" | "oficina" | "bodega"
Valores "operacion": "venta" | "renta"
Valores "condicion_terreno": "plano" | "pendiente" | "irregular" | "" (solo terrenos)

══════════════════════════════════════════════════
ACCIÓN 3: GENERAR CONTRATO DE ARRENDAMIENTO
══════════════════════════════════════════════════
Datos OBLIGATORIOS:
1. Calle del inmueble arrendado
2. Número exterior
3. Colonia
4. C.P.
5. Municipio y estado (ej: "Morelia, Michoacán")
6. Nombre completo del arrendador (dueño) — EN MAYÚSCULAS
7. Nombre completo del arrendatario (inquilino) — EN MAYÚSCULAS
8. Renta mensual (MXN)
9. Depósito en garantía (si no sabe, usa el mismo valor que la renta)
10. Fecha de inicio (día/mes/año)

[ACCION]{"tipo":"llenar_contrato","subtipo":"arrendamiento","calle_inmueble":"AV. CAMELINAS","num_ext":"123","num_int":"","colonia":"CHAPULTEPEC","cp":"58260","municipio_estado":"MORELIA, MICHOACÁN","arrendador":"SALVADOR BOLAÑOS NAVARRO","arrendatario":"GABRIELA NAVARRO PÉREZ","renta":8500,"deposito":8500,"dia_pago":5,"fecha_inicio":"2026-05-01"}[/ACCION]

dia_pago: día límite del mes para pagar (default 5). fecha_inicio en formato YYYY-MM-DD.

══════════════════════════════════════════════════
ACCIÓN 4: GENERAR PROMESA DE COMPRAVENTA
══════════════════════════════════════════════════
Datos OBLIGATORIOS:
1. Dirección del inmueble (calle y número)
2. Colonia
3. C.P.
4. Nombre del vendedor
5. Nombre del comprador
6. Precio total de venta
7. Monto de arras/enganche
8. Fecha límite para escriturar

[ACCION]{"tipo":"llenar_contrato","subtipo":"promesa","dir":"Cipres 167","colonia":"Melchor Ocampo","cp":"58160","vendedor":"JUAN PÉREZ GARCÍA","comprador":"MARÍA LÓPEZ HERNÁNDEZ","precio":2500000,"arras":250000,"fecha_limite":"2026-06-30"}[/ACCION]

fecha_limite en formato YYYY-MM-DD.

══════════════════════════════════════════════════
ACCIÓN 5: FICHA TÉCNICA DESDE EASYBROKER
══════════════════════════════════════════════════
[ACCION]{"tipo":"crear_ficha","id_easybroker":"EB-KH4322"}[/ACCION]
Si el usuario no da el ID: [ACCION]{"tipo":"navegar","modulo":"ficha"}[/ACCION]

══════════════════════════════════════════════════
ACCIÓN 6: FICHA TÉCNICA MANUAL
══════════════════════════════════════════════════
Datos mínimos: tipo, operación, precio, colonia.
[ACCION]{"tipo":"crear_ficha_manual","tipo_inmueble":"casa","operacion":"venta","precio":3500000,"colonia":"Chapultepec","ciudad":"Morelia","calle":"Av. Madero 123","recamaras":3,"banos":2,"m2_construccion":180,"m2_terreno":220,"estacionamientos":2,"descripcion":""}[/ACCION]

══════════════════════════════════════════════════
ACCIÓN 7: BUSCAR PROPIEDAD EN MIS INMUEBLES
══════════════════════════════════════════════════
[ACCION]{"tipo":"buscar_propiedad","query":"Chapultepec"}[/ACCION]

══════════════════════════════════════════════════
ACCIÓN 8: CREAR CAMPAÑA DE META ADS
══════════════════════════════════════════════════
Datos OBLIGATORIOS:
1. ¿Para qué propiedad? (nombre o descripción breve)
2. ¿Presupuesto diario en pesos? (mínimo $50)
3. Objetivo — ofrece opciones: a) Conseguir leads  b) Llevar tráfico a web  c) Reconocimiento

[ACCION]{"tipo":"confirmar_campana","nombre":"NOMBRE","objetivo":"OUTCOME_LEADS","presupuesto_diario_mxn":150,"ciudad":"Morelia","edad_min":25,"edad_max":55,"url_destino":"","texto_anuncio":""}[/ACCION]

Valores "objetivo": "OUTCOME_LEADS" | "OUTCOME_TRAFFIC" | "OUTCOME_AWARENESS"
NUNCA ejecutes sin confirmación explícita.

══════════════════════════════════════════════════
ACCIÓN 9: NAVEGAR A UN MÓDULO
══════════════════════════════════════════════════
[ACCION]{"tipo":"navegar","modulo":"isr"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"contratos"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"avm"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"props"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"ficha"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"ficha-manual"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"facebook-ads"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"contactos"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"image-cleaner"}[/ACCION]

══════════════════════════════════════════════════
ACCIÓN 10: AGREGAR CONTACTO DIRECTAMENTE (sin navegar)
══════════════════════════════════════════════════
Cuando el usuario pide agregar un contacto/prospecto/cliente, captura los datos y lánzalo directo. NO navegues. El contacto se crea en el CRM y aparece la confirmación en el chat.

Datos OBLIGATORIOS: nombre. Opcionales: telefono, email, empresa, tipo_contacto (prospecto|vendedor|comprador|arrendatario), notas.

[ACCION]{"tipo":"agregar_contacto","nombre":"María López","telefono":"4431234567","email":"maria@example.com","tipo_contacto":"prospecto","notas":"Interesada en Chapultepec, presupuesto 4M"}[/ACCION]

Ejemplo:
Usuario: "agrega a María López, su tel es 443 123 4567, le interesa una casa en Chapultepec con presupuesto de 4 millones"
Broq: "Listo, lo agrego."
[ACCION]{"tipo":"agregar_contacto","nombre":"María López","telefono":"4431234567","tipo_contacto":"prospecto","notas":"Interesada en Chapultepec, presupuesto 4M"}[/ACCION]

══════════════════════════════════════════════════
ACCIÓN 11A: CALCULAR ISR Y DESCARGAR PDF DIRECTAMENTE (preferida)
══════════════════════════════════════════════════
Cuando tengas TODOS los datos del ISR y el usuario quiere el resultado YA,
usa esta acción. El PDF se descarga directo en su dispositivo sin sacarlo
del chat. Es la acción DEFAULT para "calcular ISR" / "dame el ISR de…".

Mismos campos que `llenar_isr`, solo cambia el tipo.

[ACCION]{"tipo":"calcular_isr_directo","precio_venta":3200000,"precio_compra":1000000,"anio_venta":2026,"mes_venta":3,"anio_compra":2015,"mes_compra":1,"inmueble":"casa","exencion":"no","mejoras":0,"escrituracion":0,"comision":96000}[/ACCION]

Ejemplo:
Usuario: "calcula el ISR y mándame el PDF"
Broq: "Listo, calculando y descargando."
[ACCION]{"tipo":"calcular_isr_directo",...}[/ACCION]

══════════════════════════════════════════════════
ACCIÓN 11B: ESTIMAR VALOR Y DESCARGAR PDF DIRECTAMENTE (preferida)
══════════════════════════════════════════════════
Cuando tengas los datos para una estimación de valor y el usuario quiere el
PDF YA, usa esta acción. Busca comparables, hace el cálculo y descarga el PDF
directo en el chat. Tarda 30s–2 min porque consulta portales en vivo.

Mismos campos que `opinion_valor_web`.

[ACCION]{"tipo":"estimar_valor_directo","colonia":"Vistas Altozano","tipo_inmueble":"casa","operacion":"venta","m2_construccion":180,"m2_terreno":220,"recamaras":3,"banos":2,"ciudad":"Morelia","condicion_terreno":""}[/ACCION]

Ejemplo:
Usuario: "estima el valor de una casa de 180m² en Vistas Altozano y mándame el PDF"
Broq: "Voy a buscar comparables y prepararte el PDF. Tarda un poco."
[ACCION]{"tipo":"estimar_valor_directo",...}[/ACCION]

══════════════════════════════════════════════════
ACCIÓN 12: GENERAR Y DESCARGAR CONTRATO DIRECTAMENTE
══════════════════════════════════════════════════
Cuando ya tienes TODOS los datos obligatorios y el usuario CONFIRMA que quiere descargar el contrato, usa esta acción. El DOCX se descarga directo en su dispositivo, sin navegar.

Si faltan datos: usa "llenar_contrato" (acción 4) en su lugar — eso navega y deja el form pre-llenado para que complete.

Datos: TODOS los del contrato. subtipo: "arrendamiento" | "promesa".

[ACCION]{"tipo":"generar_contrato_directo","subtipo":"arrendamiento","datos":{...}}[/ACCION]

Ejemplo:
Usuario: "ya tengo todo, descárgame el contrato ya"
Broq: "Listo, lo genero y se descarga."
[ACCION]{"tipo":"generar_contrato_directo","subtipo":"arrendamiento","datos":{"fecha_contrato":"2026-05-21","calle_inmueble":"Av. Camelinas","num_ext_inmueble":"123","colonia_inmueble":"CHAPULTEPEC","cp_inmueble":"58260","municipio_estado_inmueble":"MORELIA, MICHOACAN","nombre_arrendador":"SALVADOR BOLAÑOS","nombre_arrendatario":"GABRIELA NAVARRO","renta_mensual":8500,"deposito_garantia":8500,"dia_pago":5,"fecha_inicio":"2026-06-01","fecha_fin":"2027-05-31"}}[/ACCION]

══════════════════════════════════════════════════
EJEMPLOS DE CONVERSACIÓN
══════════════════════════════════════════════════

EJEMPLO ISR:
Usuario: "calcula el ISR de una casa que vendí en 3.2 millones en marzo 2026, la compré en un millón en enero 2015, comisión 96 mil"
Broq: "Venta $3,200,000 en marzo 2026, compra $1,000,000 en enero 2015, comisión $96,000. ¿Usaste la exención de casa habitación en los últimos 3 años?"
Usuario: "no"
Broq: "Listo, calculando."
[ACCION]{"tipo":"llenar_isr","precio_venta":3200000,"precio_compra":1000000,"anio_venta":2026,"mes_venta":3,"anio_compra":2015,"mes_compra":1,"inmueble":"casa","exencion":"no","mejoras":0,"escrituracion":0,"comision":96000}[/ACCION]

EJEMPLO CONTRATO:
Usuario: "necesito un contrato de renta para Av. Camelinas 123, Chapultepec, Morelia, el dueño es Salvador Bolaños y el inquilino es Gabriela Navarro, renta 8,500"
Broq: "¿Cuál es el código postal?"
Usuario: "58260"
Broq: "¿Desde qué fecha empieza?"
Usuario: "1 de mayo de 2026"
Broq: "Generando contrato."
[ACCION]{"tipo":"llenar_contrato","subtipo":"arrendamiento","calle_inmueble":"AV. CAMELINAS","num_ext":"123","num_int":"","colonia":"CHAPULTEPEC","cp":"58260","municipio_estado":"MORELIA, MICHOACÁN","arrendador":"SALVADOR BOLAÑOS NAVARRO","arrendatario":"GABRIELA NAVARRO PÉREZ","renta":8500,"deposito":8500,"dia_pago":5,"fecha_inicio":"2026-05-01"}[/ACCION]

EJEMPLO EASYBROKER:
Usuario: "¿cómo conecto mi EasyBroker?"
Broq: "En EasyBroker, haz clic en tu nombre (arriba a la derecha) → Configuración de cuenta → Integraciones o API → copia tu API Key. Luego en Broquer, abre tu perfil (tus iniciales abajo a la izquierda del sidebar) → sección EasyBroker → pega la key y haz clic en Conectar. Queda listo en segundos."

EJEMPLO ASESORÍA LEGAL:
Usuario: "¿puedo rentar sin contrato escrito?"
Broq: "Técnicamente sí — el Código Civil de Michoacán permite arrendamiento verbal. Pero sin contrato escrito, si hay conflicto, la ley presume que el plazo es mensual y que no hay depósito, lo que te deja sin herramienta legal. Siempre conviene tener el contrato firmado."

Responde siempre en español. Sin markdown en respuestas conversacionales (sin **, sin #, sin listas con guiones). Usa oraciones naturales y cortas cuando el usuario habla por voz."""

class ClaudeChatRequest(BaseModel):
    messages: list
    max_tokens: int = 1200
    temperature: float = 0.7
    context: str = ""  # Módulo/pantalla activa — se inyecta al system prompt

@app.post("/chat-claude")
async def chat_claude_proxy(req: ClaudeChatRequest, request: Request):
    _uid = await get_user_id_from_token(request)
    exigir_cupo(request, _uid)
    exigir_sesion(request, _uid)
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY no configurada en el servidor")
    user_id = await get_user_id_from_token(request)

    # Construir system prompt con contexto dinámico del módulo activo
    system_content = SHAARK_SYSTEM_PROMPT
    if req.context:
        system_content += f"\n\n═══════════════════════════════════════\nCONTEXTO ACTUAL DEL USUARIO\n═══════════════════════════════════════\nEl usuario está en: {req.context}\nAdapta tu respuesta y acciones a este módulo cuando sea relevante."

    user_messages = [m for m in req.messages if m.get("role") != "system"]

    async with httpx.AsyncClient(timeout=60) as client:
        r = await client.post(
            f"{ANTHROPIC_BASE}/messages",
            headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
            },
            json={
                "model": "claude-sonnet-4-6",
                "max_tokens": req.max_tokens,
                "system": system_content,
                "messages": user_messages,
                "tools": [{"type": "web_search_20250305", "name": "web_search", "max_uses": 3}],
            }
        )
        if r.status_code != 200:
            raise HTTPException(status_code=r.status_code,
                detail=f"Error Claude: {r.text}")

        data = r.json()
        _track_anthropic(user_id, _request_modulo(request, "chat"), "/chat-claude", data,
                         modelo=data.get("model") or "claude-sonnet-4-6")
        # Extraer texto ignorando bloques tool_use (web_search)
        blocks = data.get("content", [])
        text_parts = [b.get("text", "") for b in blocks if b.get("type") == "text"]
        reply_text = "".join(text_parts).strip() or "Sin respuesta."
        return {
            "choices": [
                {"message": {"role": "assistant", "content": reply_text}}
            ]
        }


# ──────────────────────────────────────────────────────────────
# SOLICITUD DE ARRENDAMIENTO — Análisis con Claude (vision/PDF/DOCX)
# ──────────────────────────────────────────────────────────────
@app.post("/solicitud-arrendamiento/analizar")
async def analizar_solicitud_arrendamiento(
    request: Request,
    file: UploadFile = File(...),
    documentos: List[UploadFile] = File(default=[]),
):
    """
    Lee una solicitud de arrendamiento (PDF, imagen JPG/PNG/WEBP o DOCX) más
    hasta 5 documentos de respaldo opcionales (comprobantes de ingresos, escrituras
    del aval, INE, estados de cuenta, etc.) y los cruza todos con Claude Sonnet 4.6.
    Devuelve JSON estructurado con puntaje, riesgo, hallazgos y recomendaciones.
    Solicitud principal: máx 15 MB. Documentos adicionales: máx 8 MB c/u.
    Requiere usuario autenticado.
    """
    # Auth
    user_id = await get_user_id_from_token(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="Inicia sesión para usar este módulo.")
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY no configurada en el servidor.")

    # Leer archivo y validar tamaño
    content = await file.read()
    if len(content) > 15 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Archivo demasiado grande (máx 15 MB).")
    if len(content) < 100:
        raise HTTPException(status_code=400, detail="Archivo vacío o corrupto.")

    fname = (file.filename or "").lower()
    ctype = (file.content_type or "").lower()

    is_pdf = ctype == "application/pdf" or fname.endswith(".pdf")
    is_docx = "wordprocessingml" in ctype or fname.endswith(".docx")
    is_image = (
        ctype.startswith("image/")
        or any(fname.endswith(x) for x in [".jpg", ".jpeg", ".png", ".webp", ".gif"])
    )

    # System prompt: rúbrica de evaluación + formato JSON estricto
    SYSTEM_PROMPT = """Eres un perito experto en evaluación de solicitudes de arrendamiento inmobiliario en México. Analizas con el rigor de un banco o inmobiliaria seria. Detectas inconsistencias, riesgos de impago y posibles fraudes.

Envuelve tu respuesta SIEMPRE entre las etiquetas <output> y </output>. Dentro de esas etiquetas coloca ÚNICAMENTE el JSON, sin texto adicional, sin bloques de markdown, sin comentarios. Así:
<output>
{ ... tu JSON aquí ... }
</output>

La estructura del JSON debe ser:
{
  "puntaje": <entero 0-100>,
  "nivel_riesgo": "verde" | "amarillo" | "rojo",
  "veredicto_corto": "<1-2 líneas resumiendo el caso>",
  "datos_extraidos": {
    "nombre_solicitante": "<string o null>",
    "edad": "<string o null>",
    "ocupacion": "<string o null>",
    "ingresos_mensuales_mxn": <número o null>,
    "renta_solicitada_mxn": <número o null>,
    "ratio_ingreso_renta": <número o null>,
    "tiene_aval": <true | false | null>,
    "tiene_referencias": <true | false | null>
  },
  "secciones": [
    {"categoria": "Identificación", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Domicilio", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Empleo e ingresos", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Estabilidad y referencias", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Fiador o garantía", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Indicadores PLD", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Coherencia documental", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]}
  ],
  "banderas_rojas": ["..."],
  "recomendaciones": ["..."]
}

Rúbrica de puntaje:
- 90-100 (verde): completo, coherente, ratio ingreso/renta >= 3x, aval sólido con propiedad libre de gravamen
- 75-89 (verde): mayoritariamente completo, ratio 2.5-3x, mínimas faltantes
- 60-74 (amarillo): incompleto pero rescatable, ratio 2-2.5x o aval débil
- 40-59 (amarillo/rojo): faltan elementos críticos, ratio 1.5-2x, o referencias no verificables
- 0-39 (rojo): inconsistencias graves, posibles indicios de falsificación, datos críticos ausentes, ratio < 1.5x

Reglas estrictas:
1. Si no puedes extraer un dato, ponlo en null. NUNCA inventes información.
2. Calcula ratio_ingreso_renta = ingresos_mensuales_mxn / renta_solicitada_mxn cuando ambos estén presentes. Devuélvelo con 2 decimales.
3. En "secciones" SIEMPRE devuelve las 7 categorías en ese orden, aunque alguna esté "faltante".
4. estatus "faltante" = la solicitud simplemente no incluyó esa información (no es necesariamente malo, pero hay que pedirla).
5. estatus "critico" = riesgo grave detectado (no solo "falta", sino algo activamente alarmante).
6. Los "puntos" deben ser observaciones CONCRETAS, no generalidades. Cita datos específicos del documento cuando puedas.
7. "banderas_rojas" solo si hay riesgos genuinos: inconsistencias entre secciones, ratio < 2x sin aval, datos manipulados, referencias laborales sospechosas, fecha de emisión muy antigua, etc.
8. "recomendaciones" son acciones concretas que el agente debe hacer ANTES de firmar: verificar X comprobante con el patrón, confirmar Y referencia, pedir Z documento faltante, etc.
9. Indicadores PLD: revisa si hay coincidencias con criterios de actividad vulnerable de LFPIORPI (renta mensual >= 1,605 UMA = $188,282.55 MXN en 2026 obliga identificación del cliente; >= 3,210 UMA = $376,565 MXN obliga aviso al SAT)."""

    # ── Helper: convierte un UploadFile a bloque(s) de contenido para Claude ──
    async def archivo_a_bloques(uf: UploadFile, etiqueta: str, max_bytes: int = 8 * 1024 * 1024):
        """Devuelve lista de bloques content para Claude según tipo de archivo."""
        raw = await uf.read()
        if len(raw) > max_bytes or len(raw) < 50:
            return []  # omitir silenciosamente si excede límite o está vacío
        n = (uf.filename or "").lower()
        ct = (uf.content_type or "").lower()
        bloques = []
        bloques.append({"type": "text", "text": f"\n--- {etiqueta} ({uf.filename}) ---"})
        if ct == "application/pdf" or n.endswith(".pdf"):
            bloques.append({
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": base64.standard_b64encode(raw).decode("utf-8")
                }
            })
        elif "wordprocessingml" in ct or n.endswith(".docx"):
            try:
                from docx import Document as _DocxDocument
                _doc = _DocxDocument(io.BytesIO(raw))
                _parts = []
                for _p in _doc.paragraphs:
                    if _p.text and _p.text.strip():
                        _parts.append(_p.text.strip())
                for _tbl in _doc.tables:
                    for _row in _tbl.rows:
                        for _cell in _row.cells:
                            for _p in _cell.paragraphs:
                                if _p.text and _p.text.strip():
                                    _parts.append(_p.text.strip())
                _txt = "\n".join(_parts)[:10000]
                if _txt.strip():
                    bloques.append({"type": "text", "text": _txt})
            except Exception:
                pass  # omitir si no se puede leer
        elif ct.startswith("image/") or any(n.endswith(x) for x in [".jpg", ".jpeg", ".png", ".webp"]):
            _mt = "image/jpeg"
            if n.endswith(".png") or "png" in ct:
                _mt = "image/png"
            elif n.endswith(".webp") or "webp" in ct:
                _mt = "image/webp"
            bloques.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": _mt,
                    "data": base64.standard_b64encode(raw).decode("utf-8")
                }
            })
        return bloques

    # ── Construir user_content: solicitud principal ──────────────────────────
    user_content = []

    if is_pdf:
        b64 = base64.standard_b64encode(content).decode("utf-8")
        user_content.append({"type": "text", "text": "--- SOLICITUD DE ARRENDAMIENTO (documento principal) ---"})
        user_content.append({
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": b64}
        })

    elif is_docx:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            parts = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text and p.text.strip():
                                parts.append(p.text.strip())
            extracted = "\n".join(parts)[:18000]
            if not extracted.strip():
                raise HTTPException(status_code=400, detail="El DOCX no contiene texto legible.")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"No se pudo leer el DOCX: {e}")
        user_content.append({
            "type": "text",
            "text": "--- SOLICITUD DE ARRENDAMIENTO (documento principal, formato Word) ---\n\n" + extracted
        })

    elif is_image:
        media_type = "image/jpeg"
        if fname.endswith(".png") or "png" in ctype:
            media_type = "image/png"
        elif fname.endswith(".webp") or "webp" in ctype:
            media_type = "image/webp"
        elif fname.endswith(".gif") or "gif" in ctype:
            media_type = "image/gif"
        b64 = base64.standard_b64encode(content).decode("utf-8")
        user_content.append({"type": "text", "text": "--- SOLICITUD DE ARRENDAMIENTO (documento principal) ---"})
        user_content.append({
            "type": "image",
            "source": {"type": "base64", "media_type": media_type, "data": b64}
        })

    else:
        raise HTTPException(
            status_code=400,
            detail="Formato no soportado. Sube PDF, imagen (JPG/PNG/WEBP) o DOCX."
        )

    # ── Documentos adicionales (hasta 5) ─────────────────────────────────────
    docs_validos = (documentos or [])[:5]
    nombres_extra = []
    for i, doc_extra in enumerate(docs_validos, start=1):
        etiqueta = f"DOCUMENTO DE RESPALDO #{i}"
        bloques = await archivo_a_bloques(doc_extra, etiqueta)
        if bloques:
            user_content.extend(bloques)
            nombres_extra.append(doc_extra.filename or f"documento_{i}")

    # ── Instrucción final con contexto de documentos enviados ─────────────────
    if nombres_extra:
        USER_INSTRUCTION = (
            f"Se adjuntan {len(nombres_extra)} documento(s) de respaldo además de la solicitud principal: "
            + ", ".join(nombres_extra) + ".\n"
            "Cruza la información de todos los documentos entre sí:\n"
            "- Verifica que los ingresos declarados en la solicitud coincidan con los comprobantes.\n"
            "- Verifica que el aval tenga solvencia real según su escritura u otro documento.\n"
            "- Detecta inconsistencias entre lo declarado en la solicitud y lo que muestran los respaldos.\n"
            "- Menciona discrepancias específicas en la sección 'Coherencia documental' y en banderas_rojas si aplica.\n\n"
            "Devuelve tu evaluación ÚNICAMENTE dentro de etiquetas <output></output>, "
            "como se indica en el system prompt. Solo JSON entre esas etiquetas."
        )
    else:
        USER_INSTRUCTION = (
            "Analiza esta solicitud de arrendamiento. "
            "Devuelve tu evaluación ÚNICAMENTE dentro de etiquetas <output></output>, "
            "como se indica en el system prompt. Solo JSON entre esas etiquetas, nada más."
        )

    user_content.append({"type": "text", "text": USER_INSTRUCTION})

    # Llamada a Claude
    try:
        async with httpx.AsyncClient(timeout=150) as client:
            r = await client.post(
                f"{ANTHROPIC_BASE}/messages",
                headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "claude-sonnet-4-6",
                    "max_tokens": 4096,
                    "system": SYSTEM_PROMPT,
                    "messages": [{"role": "user", "content": user_content}]
                }
            )
        if r.status_code != 200:
            err_txt = (r.text or "")[:300]
            raise HTTPException(
                status_code=502,
                detail=f"Error Claude {r.status_code}: {err_txt}"
            )

        data = r.json()
        _track_anthropic(user_id, "solicitud-arr", "/solicitud-arrendamiento/analizar",
                         data, modelo=data.get("model") or "claude-sonnet-4-6")
        reply_text = ""
        try:
            reply_text = data.get("content", [{}])[0].get("text", "")
        except Exception:
            pass
        if not reply_text:
            raise HTTPException(status_code=502, detail="Claude devolvió respuesta vacía.")

        # ── Extracción robusta del JSON ──────────────────────────────────
        # Prioridad 1: contenido entre <output>...</output>
        json_str = None
        tag_match = re.search(r'<output>\s*(.*?)\s*</output>', reply_text, re.DOTALL | re.IGNORECASE)
        if tag_match:
            json_str = tag_match.group(1).strip()
        else:
            # Prioridad 2: primer bloque { ... } del texto
            brace_match = re.search(r'\{.*\}', reply_text, re.DOTALL)
            if brace_match:
                json_str = brace_match.group().strip()

        if not json_str:
            raise HTTPException(status_code=502, detail="Claude no devolvió JSON válido.")

        # Limpiar caracteres de control que vienen de PDFs (null bytes, BOM, etc.)
        # Conservamos \n \r \t que son válidos en JSON.
        json_str = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', json_str)
        # Quitar BOM si quedó al inicio
        json_str = json_str.lstrip('\ufeff')

        try:
            parsed = json.loads(json_str)
        except json.JSONDecodeError as e:
            # Segundo intento: escapar comillas dobles problemáticas dentro de valores string.
            # Reemplaza secuencias tipo :" texto "con comillas" ": con versión escapada.
            try:
                json_str2 = re.sub(
                    r'(?<=[:{,\[])\s*"((?:[^"\\]|\\.)*)"\s*(?=[,}\]:])',
                    lambda m: '"' + m.group(1).replace('"', '\\"') + '"',
                    json_str
                )
                parsed = json.loads(json_str2)
            except Exception:
                raise HTTPException(
                    status_code=502,
                    detail=f"JSON inválido de Claude: {str(e)[:120]}"
                )

        # Validación ligera del shape
        if "puntaje" not in parsed or "nivel_riesgo" not in parsed:
            raise HTTPException(status_code=502, detail="Respuesta sin estructura esperada.")

        # Asegurar que datos_extraidos y secciones existan (aunque vacías)
        parsed.setdefault("datos_extraidos", {})
        parsed.setdefault("secciones", [])
        parsed.setdefault("banderas_rojas", [])
        parsed.setdefault("recomendaciones", [])
        parsed.setdefault("veredicto_corto", "")

        return parsed

    except HTTPException:
        raise
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="El análisis tardó demasiado. Intenta de nuevo.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando: {str(e)[:200]}")


# ════════════════════════════════════════════════════════════════
# IMPORTACIÓN MASIVA DESDE EASYBROKER
# Trae TODAS las propiedades del agente desde su cuenta de EasyBroker
# y las inserta en Supabase (tabla propiedades) bajo SU user_id.
# Deduplicación por eb_public_id: si ya existe, la salta.
# ════════════════════════════════════════════════════════════════

@app.post("/easybroker/import-all")
async def easybroker_import_all(request: Request):
    """
    Importa propiedades PUBLICADAS del agente desde su cuenta de EasyBroker
    a Mis Inmuebles. Upsert por eb_public_id: si ya existe, actualiza datos
    de EB pero PRESERVA notas internas y estatus que el usuario haya cambiado.

    Optimizaciones:
    - Filtra solo published con search[statuses][]=published
    - Procesa detalles en paralelo (lotes de 10)
    - Inserta en lotes a Supabase (1 POST por lote, no 1 por propiedad)
    - Preserva notas y estatus del usuario en filas existentes
    """
    user_id = await get_user_id_from_token(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="Tu sesión expiró. Vuelve a iniciar sesión.")
    user_key = await get_eb_key_for_user(user_id)
    if not user_key:
        raise HTTPException(status_code=400, detail="Configura tu API key de EasyBroker en Perfil → Integración EasyBroker antes de importar.")
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=500, detail="Supabase no está configurado en el servidor.")

    # ─── Estatus elegidos por el usuario ───
    # Body opcional: {"statuses": ["published", "sold", ...]}
    try:
        body_imp = await request.json()
    except Exception:
        body_imp = {}
    # Con fotos_diferidas=true NO se lanza la copia de fotos al terminar.
    # La migración completa lo usa para que la copia (pesada) no compita con
    # los pasos de contactos e historial en el mismo worker.
    fotos_diferidas = bool((body_imp or {}).get("fotos_diferidas"))
    pedidos = (body_imp or {}).get("statuses")
    if isinstance(pedidos, str):
        pedidos = [pedidos]
    if isinstance(pedidos, list):
        statuses_elegidos = [s for s in _EB_STATUS_MAP if s in pedidos]
    else:
        statuses_elegidos = []
    if not statuses_elegidos:
        statuses_elegidos = list(_EB_STATUS_DEFAULT)

    sb_headers = {
        "apikey": SUPABASE_SERVICE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        "Content-Type": "application/json",
    }

    # ─── Paso 1: leer filas existentes del usuario (para preservar notas/estatus) ───
    existentes_por_eb_id = {}  # eb_public_id → {notas, estatus}
    try:
        try:
            filas_existentes = await get_rows(
                "propiedades",
                {"user_id": f"eq.{user_id}",
                 "eb_public_id": "not.is.null",
                 "select": "eb_public_id,notas,estatus"},
                timeout=15,
            )
        except httpx.HTTPStatusError:
            filas_existentes = []
        for row in filas_existentes:
            eb_id = row.get("eb_public_id")
            if eb_id:
                existentes_por_eb_id[eb_id] = {
                    "notas":   row.get("notas"),
                    "estatus": row.get("estatus"),
                }
    except Exception as e:
        print(f"[import-all] Error leyendo existentes: {e}")

    # ─── Paso 2: paginar el listado de EasyBroker, un estatus a la vez ───
    # IMPORTANTE: EasyBroker NO incluye el estatus dentro de cada propiedad,
    # ni en el listado ni en el detalle. La única forma de saber de qué estatus
    # es una propiedad es preguntarle por ese estatus y etiquetar lo que venga.
    # Por eso paginamos un estatus a la vez. (Verificado con /easybroker/diagnostico)
    estatus_por_pid = {}     # public_id → estatus Broquer
    conteo_por_estatus = {}  # estatus EB → cuántas llegaron
    ids_published = []       # orden de llegada (nombre histórico, se conserva)
    limite_alcanzado = False
    descartadas_estatus = 0  # repetidas entre estatus (ya contadas en otro)
    for s in statuses_elegidos:
        conteo_por_estatus[s] = 0

    async with httpx.AsyncClient(timeout=30) as client:
        for eb_status in statuses_elegidos:
            if limite_alcanzado:
                break
            brokr_status = _EB_STATUS_MAP[eb_status]
            pagina = 1
            while pagina <= 400:  # tope duro de seguridad
                r = await _eb_get_reintentos(
                    client,
                    f"{EB_BASE}/properties",
                    eb_headers(user_key),
                    [("limit", 50), ("page", pagina),
                     ("search[statuses][]", eb_status)],
                    timeout=30.0,
                )
                if r is None:
                    break
                if r.status_code == 401:
                    raise HTTPException(status_code=401, detail="Tu API key de EasyBroker fue rechazada. Reconéctala en Perfil.")
                if r.status_code != 200:
                    break
                data = r.json()
                content = data.get("content", []) or []
                if not content:
                    break
                for p in content:
                    if len(ids_published) >= _EB_LIMITE_PROPIEDADES:
                        limite_alcanzado = True
                        break
                    pid = p.get("public_id")
                    if not pid:
                        continue
                    if pid in estatus_por_pid:
                        descartadas_estatus += 1
                        continue
                    estatus_por_pid[pid] = brokr_status
                    conteo_por_estatus[eb_status] = conteo_por_estatus.get(eb_status, 0) + 1
                    ids_published.append(pid)
                if limite_alcanzado:
                    break
                if not data.get("pagination", {}).get("next_page"):
                    break
                pagina += 1

    total_eb = len(ids_published)

    # ─── Paso 3: traer detalle de TODAS las published en paralelo (lotes de 10) ───
    # Aún las que ya existen las re-procesamos para que el upsert actualice precio,
    # fotos, descripción, amenidades, etc. (Decisión D2).
    errores: list = []
    inmuebles_listos: list = []

    # La empresa comparte UNA cuenta de EasyBroker. Sin esto, cada agente que
    # importara crearía su propia copia del mismo inventario.
    org_id_import = await get_org_id_for_user(user_id)
    if not org_id_import:
        raise HTTPException(status_code=403, detail="Tu cuenta no está configurada. Contacta a soporte.")

    async def fetch_one(client: httpx.AsyncClient, pid: str):
        try:
            rd = await _eb_get_reintentos(
                client,
                f"{EB_BASE}/properties/{pid}",
                eb_headers(user_key),
                timeout=20.0,
            )
            if rd is None:
                return ("err", {"id": pid, "error": "EasyBroker no respondió tras varios intentos"})
            if rd.status_code != 200:
                return ("err", {"id": pid, "error": f"EB status {rd.status_code}"})
            prop_full = rd.json()
            inmueble = _eb_to_brokr(prop_full, user_id)
            inmueble["org_id"] = org_id_import
            # EasyBroker no manda el estatus dentro de la propiedad. Usamos el
            # estatus por el que preguntamos para traerla.
            eb_estatus = estatus_por_pid.get(pid)
            if eb_estatus:
                inmueble["estatus"] = eb_estatus
            # Preservar notas y estatus del usuario si la fila ya existe
            prev = existentes_por_eb_id.get(pid)
            if prev:
                if prev.get("notas"):
                    inmueble["notas"] = prev["notas"]
                if prev.get("estatus"):
                    inmueble["estatus"] = prev["estatus"]
            return ("ok", inmueble)
        except Exception as e:
            return ("err", {"id": pid, "error": str(e)[:120]})

    BATCH = _EB_LOTE
    lotes_fallidos_seguidos = 0
    async with httpx.AsyncClient(timeout=30) as client:
        for i in range(0, len(ids_published), BATCH):
            chunk = ids_published[i:i+BATCH]
            _prog(user_id, f"propiedades {min(i + BATCH, len(ids_published))} de {len(ids_published)}")
            inicio_lote = time.monotonic()
            results = await asyncio.gather(*[fetch_one(client, pid) for pid in chunk])
            # Mantener el ritmo por debajo del límite de EasyBroker: si el lote
            # tardó menos que la pausa mínima, esperamos la diferencia.
            resto = _EB_PAUSA_LOTE - (time.monotonic() - inicio_lote)
            if resto > 0 and i + BATCH < len(ids_published):
                await asyncio.sleep(resto)
            fallos_lote = 0
            for status, payload in results:
                if status == "ok":
                    inmuebles_listos.append(payload)
                else:
                    errores.append(payload)
                    fallos_lote += 1
            # Cortacircuito: si EasyBroker rechaza TODO durante varios lotes
            # seguidos (429 sostenido), no tiene caso moler reintentos media
            # hora. Se aborta con mensaje claro.
            lotes_fallidos_seguidos = (lotes_fallidos_seguidos + 1
                                       if fallos_lote == len(chunk) else 0)
            if lotes_fallidos_seguidos >= 4:
                raise HTTPException(status_code=429, detail="EasyBroker está limitando las peticiones de tu cuenta (429 sostenido). Espera 10-15 minutos y vuelve a correr la migración: lo ya importado no se pierde ni se duplica.")

    # ─── Paso 4: UPSERT en lotes a Supabase (50 por POST) ───
    # Necesita el índice único (user_id, eb_public_id) en Supabase para que
    # on_conflict funcione.
    upserted = 0
    UPSERT_BATCH = 50
    async with httpx.AsyncClient(timeout=60) as client:
        for i in range(0, len(inmuebles_listos), UPSERT_BATCH):
            chunk = inmuebles_listos[i:i+UPSERT_BATCH]
            ultimo_fallo = "sin respuesta"
            guardado = False
            for intento in range(3):
                try:
                    await upsert_rows(
                        "propiedades",
                        chunk,
                        conflict="org_id,eb_public_id",
                        prefer="resolution=merge-duplicates,return=minimal",
                        timeout=60,
                        accepted_statuses=(200, 201, 204),
                    )
                    upserted += len(chunk)
                    guardado = True
                    break
                except httpx.HTTPStatusError as e:
                    ultimo_fallo = f"Supabase {e.response.status_code}: {e.response.text[:200]}"
                except Exception as e:
                    ultimo_fallo = str(e)[:200]
                await asyncio.sleep(1.5 * (2 ** intento))
            if not guardado:
                errores.append({
                    "id": f"lote_{i // UPSERT_BATCH}",
                    "error": ultimo_fallo
                })

    nuevas      = sum(1 for inm in inmuebles_listos if inm["eb_public_id"] not in existentes_por_eb_id)
    actualizadas = upserted - nuevas if upserted >= nuevas else 0

    # Guardar las fotos en Broquer, solo, sin que el usuario espere ni deje
    # la pestaña abierta. Si ya hay un proceso corriendo para esta empresa,
    # el propio trabajador se ignora a sí mismo.
    fotos_lanzado = False
    if org_id_import and upserted and not fotos_diferidas:
        try:
            asyncio.create_task(_migrar_fotos_org(org_id_import))
            fotos_lanzado = True
        except Exception as e:
            print(f"[import-all] No se pudo lanzar el guardado de fotos: {e}")

    return {
        "total_easybroker": total_eb,
        "importadas":       nuevas,           # nuevas filas creadas
        "actualizadas":     actualizadas,     # ya existían y se actualizaron
        "ya_existian":      actualizadas,     # backward-compat con frontend viejo
        "por_estatus":      conteo_por_estatus,  # cuántas de cada estatus EB
        "statuses":         statuses_elegidos,   # estatus que se importaron
        "descartadas":      descartadas_estatus, # EB las mandó pero no se pidieron
        "limite":           _EB_LIMITE_PROPIEDADES,
        "limite_alcanzado": limite_alcanzado,
        "fotos_en_proceso": fotos_lanzado,
        "errores":          errores
    }


# Borrado masivo de propiedades y contactos.
from routers.bulk_delete import router as bulk_delete_router
app.include_router(bulk_delete_router)

app.include_router(admin_usage_router)
app.include_router(account_delete_router)
app.include_router(avm_legacy_router)

app.include_router(avm_claude_router)

app.include_router(avm_websearch_router)

app.include_router(facebook_connection_read_router)

app.include_router(facebook_pages_router)

app.include_router(facebook_select_page_router)

app.include_router(facebook_select_ad_account_router)

app.include_router(facebook_encrypt_tokens_router)

app.include_router(facebook_disconnect_router)

app.include_router(facebook_ad_accounts_router)

app.include_router(facebook_city_search_router)

app.include_router(facebook_campaigns_router)

app.include_router(facebook_insights_read_router)

app.include_router(facebook_campaign_review_router)

app.include_router(facebook_leadgen_verify_router)

app.include_router(facebook_leadgen_status_router)

app.include_router(facebook_leadgen_subscribe_router)

app.include_router(facebook_leadgen_webhook_router)

app.include_router(facebook_page_posts_router)

app.include_router(facebook_audiences_read_router)

app.include_router(facebook_oauth_callback_router)

app.include_router(facebook_publish_router)

app.include_router(facebook_publish_property_router)

app.include_router(facebook_save_page_router)

app.include_router(facebook_ad_description_router)

app.include_router(facebook_campaign_toggle_router)




# ────────────────────────────────────────────
# COLONIAS AUTOCOMPLETE
# ────────────────────────────────────────────
# ────────────────────────────────────────────
# AVM — HELPERS
# ────────────────────────────────────────────







# ────────────────────────────────────────────
# HEDONIC MODEL
# ────────────────────────────────────────────


# ────────────────────────────────────────────
# AVM ENDPOINT
# ────────────────────────────────────────────




# ────────────────────────────────────────────
# AVM — CLAUDE AI OPINION DE VALOR
# ────────────────────────────────────────────




# ────────────────────────────────────────────
# AVM — OPINIÓN DE VALOR CON INVESTIGACIÓN CONTROLADA DE COMPARABLES
# ────────────────────────────────────────────



# ────────────────────────────────────────────
# AVM — PDF DE OPINIÓN DE VALOR
# ────────────────────────────────────────────

@app.post("/avm-pdf")
async def generar_avm_pdf(p: dict):
    """Recibe el resultado del AVM websearch y genera un PDF profesional con Playwright.

    Sistema de diseño: los mismos tokens de brokr-theme.css (navy, azul,
    Manrope, radios, sombras) que usa el resto de Broquer — para que este
    documento se sienta hermano de la Ficha técnica y del ISR, no un
    invitado con otra identidad visual.
    """
    from playwright.async_api import async_playwright

    resultado = p.get("resultado", {})
    agente = p.get("agente", "Agente Broquer")

    if not resultado:
        raise HTTPException(status_code=400, detail="Resultado vacío")

    def fmt_mx(n):
        try:
            return "${:,.0f}".format(float(n))
        except Exception:
            return str(n)

    def _esc(s):
        return (str(s) if s is not None else "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace('"',"&quot;")

    # Comparables
    comps_html = ""
    for c in resultado.get("comparables", []):
        fuente = c.get("fuente","—") or "—"
        url = c.get("url","") or ""
        src_cell = (
            f'<a href="{_esc(url)}" target="_blank" rel="noopener" class="link">{_esc(fuente)}</a>'
            if url else _esc(fuente)
        )
        comps_html += f"""
        <tr>
          <td>{_esc(c.get('descripcion','—'))}</td>
          <td class="num">{_esc(c.get('superficie_m2','—'))} m²</td>
          <td class="num">{fmt_mx(c.get('precio',0))}</td>
          <td class="num">{fmt_mx(c.get('precio_m2',0))}/m²</td>
          <td class="src">{src_cell}</td>
        </tr>"""

    # Factores de ajuste — badge con punto, mismo patrón que .bk-badge del app
    factores_html = ""
    for f in resultado.get("factores_ajuste", []):
        imp = f.get("impacto", "neutro")
        badge_cls = "badge--success" if imp == "positivo" else "badge--danger" if imp == "negativo" else "badge--mute"
        etiqueta = "Favorable" if imp == "positivo" else "Desfavorable" if imp == "negativo" else "Neutro"
        factores_html += f"""
        <tr>
          <td>
            <div class="factor-nombre">{_esc(f.get('factor','—'))}</div>
            <span class="badge {badge_cls}"><span class="dot"></span>{etiqueta}</span>
          </td>
          <td class="factor-desc">{_esc(f.get('descripcion','—'))}</td>
        </tr>"""

    recs_html = "".join(f"<li>{_esc(r)}</li>" for r in resultado.get("recomendaciones", []))

    m2c = resultado.get("m2_construccion", 0)
    m2t = resultado.get("m2_terreno", 0)
    sup_parts = []
    if m2t: sup_parts.append(f"{m2t} m² terreno")
    if m2c: sup_parts.append(f"{m2c} m² construcción")
    superficie_str = " · ".join(sup_parts) if sup_parts else "—"

    fecha_hoy = resultado.get("fecha", time.strftime("%d/%m/%Y"))
    operacion = (resultado.get('operacion','venta') or 'venta').capitalize()

    # Tokens desde brokr-theme.css. Radios propios del documento.
    _AVM_TOKENS = theme_css_for_pdf(
        "--r-xs:4px; --r-sm:8px; --r:14px; --r-lg:28px; --r-pill:999px;"
    )
    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8"/>
<style>
{_AVM_TOKENS}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: var(--font-sans); color: var(--ink); background: var(--paper); font-size: 13px; line-height: 1.55; -webkit-font-smoothing: antialiased; letter-spacing: -0.01em; }}
  .page {{ padding: 48px 52px 40px; max-width: 780px; margin: 0 auto; }}

  /* ── Encabezado de documento ── */
  .doc-head {{ display: flex; justify-content: space-between; align-items: flex-end; padding-bottom: 20px; border-bottom: 1px solid var(--line); margin-bottom: 28px; }}
  .doc-head__brand {{ font-size: 15px; font-weight: 700; color: var(--sky-navy); letter-spacing: -0.01em; }}
  .doc-head__title {{ font-size: 12px; color: var(--mute); margin-top: 2px; }}
  .doc-head__date {{ font-size: 11px; color: var(--mute); }}

  /* ── Bloque de valor — tarjeta navy, no negro genérico ── */
  .valor-card {{
    background: linear-gradient(155deg, var(--sky-navy), var(--sky-navy-mid));
    border-radius: var(--r-lg);
    padding: 26px 28px 22px;
    margin-bottom: 22px;
    -webkit-print-color-adjust: exact; print-color-adjust: exact;
  }}
  .valor-lbl {{ font-size: 11px; color: rgba(255,255,255,.65); font-weight: 600; letter-spacing: 0.02em; margin-bottom: 6px; }}
  .valor-num {{ font-family: var(--font-sans); font-size: 34px; font-weight: 700; color: #fff; line-height: 1.05; letter-spacing: -0.02em; }}
  .valor-meta {{ display: grid; grid-template-columns: repeat(4,1fr); gap: 18px; margin-top: 20px; padding-top: 18px; border-top: 1px solid rgba(255,255,255,.14); }}
  .meta-item .meta-lbl {{ font-size: 10px; color: rgba(255,255,255,.55); font-weight: 600; letter-spacing: 0.02em; margin-bottom: 4px; }}
  .meta-item .meta-val {{ font-size: 13px; font-weight: 700; color: #fff; letter-spacing: -0.005em; }}

  /* ── Secciones ── */
  .seccion {{ margin-bottom: 26px; }}
  .sec-titulo {{ font-size: 11px; font-weight: 700; color: var(--mute); letter-spacing: 0.02em; margin-bottom: 12px; }}
  .resumen {{ font-size: 12.5px; color: var(--ink-2); line-height: 1.7; text-align: justify; }}

  /* ── Badge con punto — idéntico a .bk-badge del app ── */
  .badge {{
    display: inline-flex; align-items: center; gap: 5px;
    padding: 3px 9px; border-radius: var(--r-pill);
    font-size: 11px; font-weight: 700; letter-spacing: 0.02em;
    background: var(--paper-2); color: var(--mute);
  }}
  .badge .dot {{ width: 6px; height: 6px; border-radius: 50%; background: currentColor; }}
  .badge--success {{ background: var(--success-soft); color: var(--success); }}
  .badge--danger  {{ background: var(--danger-soft);  color: var(--danger); }}
  .badge--mute    {{ background: var(--paper-2);       color: var(--mute); }}

  /* ── Tablas ── */
  table {{ width: 100%; border-collapse: collapse; font-size: 12px; }}
  th {{ font-weight: 700; color: var(--mute); text-align: left; padding: 8px 6px; border-bottom: 1px solid var(--line-2); font-size: 10px; letter-spacing: 0.02em; }}
  td {{ padding: 12px 6px; border-bottom: 1px solid var(--line); color: var(--ink); vertical-align: top; }}
  td.num {{ text-align: right; font-weight: 700; font-variant-numeric: tabular-nums; color: var(--ink); }}
  .link {{ color: var(--forest); text-decoration: underline; }}
  tr:last-child td {{ border-bottom: none; }}

  .factor-nombre {{ font-weight: 700; font-size: 12.5px; margin-bottom: 5px; }}
  .factor-desc {{ color: var(--mute); font-size: 11.5px; line-height: 1.5; }}

  .recs {{ padding-left: 18px; }}
  .recs li {{ font-size: 12.5px; color: var(--ink-2); line-height: 1.7; margin-bottom: 4px; }}

  .footer {{ margin-top: 40px; padding-top: 16px; border-top: 1px solid var(--line); text-align: center; font-size: 10px; color: var(--mute-2); letter-spacing: 0.02em; }}
</style>
</head>
<body>
<div class="page">

  <div class="doc-head">
    <div>
      <div class="doc-head__brand">Broquer</div>
      <div class="doc-head__title">Estimación de valor</div>
    </div>
    <div class="doc-head__date">{fecha_hoy}</div>
  </div>

  <div class="valor-card">
    <div class="valor-lbl">Valor estimado</div>
    <div class="valor-num">{fmt_mx(resultado.get('valor_estimado',0))}</div>
    <div class="valor-meta">
      <div class="meta-item">
        <div class="meta-lbl">Inmueble</div>
        <div class="meta-val">{_esc(resultado.get('tipo_inmueble','—'))}</div>
      </div>
      <div class="meta-item">
        <div class="meta-lbl">Superficie</div>
        <div class="meta-val">{_esc(superficie_str)}</div>
      </div>
      <div class="meta-item">
        <div class="meta-lbl">Ubicación</div>
        <div class="meta-val">{_esc(resultado.get('colonia','—'))}, {_esc(resultado.get('ciudad','Morelia'))}</div>
      </div>
      <div class="meta-item">
        <div class="meta-lbl">Operación</div>
        <div class="meta-val">{_esc(operacion)}</div>
      </div>
    </div>
  </div>

  <div class="seccion">
    <div class="sec-titulo">Análisis</div>
    <div class="resumen">{_esc(resultado.get('resumen_ejecutivo','—'))}</div>
  </div>

  <div class="seccion">
    <div class="sec-titulo">Comparables de mercado</div>
    <table>
      <thead>
        <tr>
          <th>Propiedad</th>
          <th style="text-align:right">Superficie</th>
          <th style="text-align:right">Precio</th>
          <th style="text-align:right">$/m²</th>
          <th>Fuente</th>
        </tr>
      </thead>
      <tbody>{comps_html}</tbody>
    </table>
  </div>

  {"" if not factores_html else f'''
  <div class="seccion">
    <div class="sec-titulo">Factores de ajuste</div>
    <table>
      <tbody>{factores_html}</tbody>
    </table>
  </div>
  '''}

  {"" if not recs_html else f'''
  <div class="seccion">
    <div class="sec-titulo">Recomendaciones</div>
    <ul class="recs">{recs_html}</ul>
  </div>
  '''}

  <div class="footer">Powered by Broquer</div>

</div>
</body>
</html>"""
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(args=["--no-sandbox", "--disable-dev-shm-usage"])
        page = await browser.new_page()
        await page.set_content(html, wait_until="domcontentloaded")
        await page.wait_for_timeout(400)
        pdf_bytes = await page.pdf(
            format="A4",
            print_background=True,
            margin={"top": "10mm", "right": "10mm", "bottom": "10mm", "left": "10mm"}
        )
        await browser.close()

    token = str(_uuid.uuid4()).replace("-", "")[:16]
    colonia_slug = resultado.get("colonia", "propiedad").replace(" ", "_")[:20]
    filename = f"Estimacion_Valor_{colonia_slug}_{time.strftime('%Y%m%d')}.pdf"
    _pdf_store[token] = (pdf_bytes, filename)
    if len(_pdf_store) > 50:
        oldest = list(_pdf_store.keys())[0]
        del _pdf_store[oldest]

    from fastapi.responses import JSONResponse
    return JSONResponse({"token": token, "filename": filename})


# ────────────────────────────────────────────
# CONTRATOS
# ────────────────────────────────────────────
import os

# ── PDF GENERATION ──────────────────────────────────────────────
from playwright.async_api import async_playwright
import base64, asyncio
from pydantic import BaseModel
from typing import List, Optional

class FotoItem(BaseModel):
    url: Optional[str] = None
    original: Optional[str] = None

class PropData(BaseModel):
    id: Optional[str] = None
    public_id: Optional[str] = None
    title: Optional[str] = None
    property_type: Optional[str] = None
    description: Optional[str] = None
    operations: Optional[list] = None
    location: Optional[dict] = None
    address: Optional[str] = None
    bedrooms: Optional[float] = None
    bathrooms: Optional[float] = None
    half_bathrooms: Optional[float] = None
    construction_size: Optional[float] = None
    lot_size: Optional[float] = None
    parking_spaces: Optional[float] = None
    floors: Optional[float] = None
    age: Optional[float] = None
    amenities: Optional[list] = None
    property_images: Optional[list] = None
    status: Optional[str] = None


def build_ficha_html(p: dict, images_b64: dict) -> str:
    """Plantilla editorial Broquer para la ficha técnica en PDF — edición Sky.
    Portada con tarjeta flotante sobre la foto, franja de specs con
    iconografía propia, galería en cuadrícula, características agrupadas
    por categoría, y footer de marca "Powered by Broquer" en cada página.
    """
    import re as _re
    id_prop  = p.get("public_id") or p.get("id") or ""
    titulo_base = p.get("title") or p.get("property_type") or "Propiedad"
    ops      = p.get("operations") or []
    sale_op   = next((o for o in ops if o.get("type") == "sale"), None)
    rental_op = next((o for o in ops if o.get("type") == "rental"), None)
    if not sale_op and not rental_op and ops:
        sale_op = ops[0]  # fallback: operación sin type explícito

    def fmt_money(op):
        if not op or not op.get("amount"):
            return None
        monto  = op.get("amount", 0)
        moneda = op.get("currency", "MXN")
        base = "${:,.0f}".format(monto)
        return base if moneda == "MXN" else base + " " + moneda

    es_venta_renta = bool(sale_op and rental_op)
    precio_venta = fmt_money(sale_op)
    precio_renta = fmt_money(rental_op)
    precio_principal = precio_venta or precio_renta or "—"
    if es_venta_renta:
        tipo_op = "Venta y renta"
    elif rental_op:
        tipo_op = "Renta"
    else:
        tipo_op = "Venta"

    loc      = p.get("location") or {}
    colonia  = (loc.get("name") or "").strip()
    ciudad   = (loc.get("city") or "").strip()
    direccion= (p.get("address") or "").strip()
    ubicacion= ", ".join(filter(None, [colonia, ciudad])) or direccion or "—"

    rec      = p.get("bedrooms")
    ban      = p.get("bathrooms")
    mban     = p.get("half_bathrooms")
    m2c      = p.get("construction_size")
    m2t      = p.get("lot_size")
    parking  = p.get("parking_spaces")
    niveles  = p.get("floors")
    anio     = p.get("age")
    desc     = (p.get("description") or "").replace("<br>", " ").replace("<br/>", " ")
    desc     = _re.sub(r"<[^>]+>", "", desc).strip()
    fotos    = p.get("property_images") or []
    amenids  = p.get("amenities") or []
    tipo_inmueble = (p.get("property_type") or "").strip()
    titulo   = titulo_base

    def asset_data_uri(filename: str, mime: str = "image/png") -> str:
        try:
            with open(filename, "rb") as fh:
                return f"data:{mime};base64," + base64.b64encode(fh.read()).decode()
        except Exception:
            return ""

    logo_white = asset_data_uri("logotipo-white.png")

    def fmt_m2(n):
        if not n:
            return None
        s = "{:,.2f}".format(n).rstrip("0").rstrip(".")
        return s + " m²"

    # ── Iconografía propia (línea 1.5px, redondeada, grid 24×24) ──
    ICO = {
        "bed":     '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M3 18v-6a3 3 0 013-3h12a3 3 0 013 3v6"/><path d="M3 18h18M3 18v2m18-2v2"/><path d="M7 12V9a1 1 0 011-1h3a1 1 0 011 1v3"/></svg>',
        "bath":    '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M5 12V6.5A2.5 2.5 0 017.5 4a2.5 2.5 0 012.5 2.5"/><path d="M3 12h18v2a5 5 0 01-5 5H8a5 5 0 01-5-5v-2z"/><path d="M6 19v2m12-2v2"/></svg>',
        "toilet":  '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M7 3.5h6a1 1 0 011 1V8H6V4.5a1 1 0 011-1z"/><path d="M5.5 8h9a2 2 0 012 2c0 6-3 10.5-6.5 10.5S3.5 16 3.5 10a2 2 0 012-2z"/></svg>',
        "area":    '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M4 9V4h5M15 4h5v9M20 15v5h-5M9 20H4v-5"/></svg>',
        "land":    '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M9 3L3 5.5v15L9 18l6 3 6-2.5v-15L15 6 9 3z"/><path d="M9 3v15M15 6v15"/></svg>',
        "parking": '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M5 11l1.4-4.4A2 2 0 018.3 5h7.4a2 2 0 011.9 1.6L19 11"/><path d="M5 11h14a1 1 0 011 1v4a1 1 0 01-1 1h-1a1 1 0 01-1-1v-1H7v1a1 1 0 01-1 1H5a1 1 0 01-1-1v-4a1 1 0 011-1z"/><circle cx="7.5" cy="16.5" r="1.3"/><circle cx="16.5" cy="16.5" r="1.3"/></svg>',
        "levels":  '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M12 2.5l8.5 4.5-8.5 4.5-8.5-4.5L12 2.5z"/><path d="M3.5 12l8.5 4.5 8.5-4.5"/><path d="M3.5 16.5L12 21l8.5-4.5"/></svg>',
        "calendar":'<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><rect x="3.5" y="5" width="17" height="15.5" rx="2"/><path d="M16 3v4M8 3v4M3.5 10h17"/></svg>',
        "tag":     '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M11.7 2.6a1.8 1.8 0 00-1.3-.5H4.3A1.8 1.8 0 002.5 3.9v6.1c0 .5.2.9.5 1.3l8 8a2.2 2.2 0 003 0l6-6a2.2 2.2 0 000-3.1l-8-8z"/><circle cx="7" cy="7.2" r="1.4"/></svg>',
        "pin":     '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M12 21.3S5.5 14.8 5.5 9.8a6.5 6.5 0 0113 0c0 5-6.5 11.5-6.5 11.5z"/><circle cx="12" cy="9.8" r="2.4"/></svg>',
        "route":   '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><circle cx="5" cy="18" r="2"/><circle cx="19" cy="6" r="2"/><path d="M7 18h7a4 4 0 004-4V9"/></svg>',
        "home":    '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M3.5 11.2L12 4l8.5 7.2"/><path d="M5.5 9.8v9.7a1 1 0 001 1H9v-6h6v6h2.5a1 1 0 001-1V9.8"/></svg>',
        "swap":    '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M4 8h13l-3.5-3.5M20 16H7l3.5 3.5"/></svg>',
        "photo":   '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="4.5" width="18" height="15" rx="2"/><circle cx="8.5" cy="10" r="1.6"/><path d="M21 15.5l-5.2-5.2a1.5 1.5 0 00-2.1 0L5 19"/></svg>',
        "sparkles":'<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M12 3l1.6 5.4L19 10l-5.4 1.6L12 17l-1.6-5.4L5 10l5.4-1.6L12 3z"/><path d="M19 15l.7 2.3 2.3.7-2.3.7-.7 2.3-.7-2.3-2.3-.7 2.3-.7.7-2.3z"/></svg>',
        "list":    '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M9 6h11M9 12h11M9 18h11"/><path d="M4.5 6h.01M4.5 12h.01M4.5 18h.01"/></svg>',
    }

    # ── Specs de portada (hasta 6) ──
    specs = []
    if rec:    specs.append((ICO["bed"], str(int(rec)) if float(rec).is_integer() else str(rec), "Recámaras"))
    if ban:    specs.append((ICO["bath"], str(int(ban)) if float(ban).is_integer() else str(ban), "Baños"))
    if mban:   specs.append((ICO["toilet"], str(int(mban)) if float(mban).is_integer() else str(mban), "Medios baños"))
    def fmt_num(n):
        if not n:
            return None
        s = "{:,.2f}".format(n).rstrip("0").rstrip(".")
        return s

    if m2c:    specs.append((ICO["area"], fmt_num(m2c), "m² const."))
    if m2t:    specs.append((ICO["land"], fmt_num(m2t), "m² terreno"))
    if parking and len(specs) < 6: specs.append((ICO["parking"], str(int(parking)) if float(parking).is_integer() else str(parking), "Estac."))
    if niveles and len(specs) < 6: specs.append((ICO["levels"], str(int(niveles)) if float(niveles).is_integer() else str(niveles), "Niveles"))
    specs = specs[:6]

    specs_items = "".join(
        '<div class="spec-item"><div class="spec-ico">{}</div><div class="spec-val">{}</div><div class="spec-lbl">{}</div></div>'.format(i, v, l)
        for i, v, l in specs
    )
    specs_html = '<div class="cover-specs" style="--spec-cols:{}">{}</div>'.format(len(specs), specs_items) if specs_items else ""

    foto_urls = [f.get("url") or f.get("original") or "" for f in fotos if f]
    foto_urls = [u for u in foto_urls if u]
    hero_src  = images_b64.get(foto_urls[0], foto_urls[0]) if foto_urls else ""
    hero_html = '<img class="cover-hero" src="{}" alt="portada"/>'.format(hero_src) if hero_src else '<div class="cover-hero-placeholder">{}</div>'.format(ICO["home"])
    total_fotos = len(foto_urls)
    photocount_html = ''
    if total_fotos:
        photocount_html = '<div class="cover-photocount">{}{} foto{}</div>'.format(ICO["photo"], total_fotos, "" if total_fotos == 1 else "s")
    brandmark_html = '<div class="cover-brandmark"><img src="{}" alt="Broquer"/></div>'.format(logo_white) if logo_white else '<div class="cover-brandmark"><strong style="color:#fff">Broquer</strong></div>'

    def footer(page_num, total_pages):
        logo = '<img src="{}" alt="Broquer"/>'.format(logo_white) if logo_white else '<strong>Broquer</strong>'
        id_html = '<span class="ft-id">{}</span>'.format(id_prop) if id_prop else ''
        return (
            '<div class="ficha-footer">'
            '<div class="ft-brand">{}<span>Powered by Broquer</span></div>'
            '<div class="ft-meta">{}<span>{} / {}</span></div>'
            '</div>'
        ).format(logo, id_html, page_num, total_pages)

    precio_sec_html = ""
    if es_venta_renta and precio_renta:
        precio_sec_html = '<div class="cover-precio-sec">También disponible en renta: <b>{}/mes</b></div>'.format(precio_renta)

    cover_content = (
        '<div class="cover-hero-wrap">{}{}{}</div>'
        '<div class="cover-card">'
        '<div class="cover-card-top">'
        '<div class="cover-precio-block">'
        '<div class="cover-badge">{}</div>'
        '<div class="cover-precio">{}</div>{}'
        '</div>'
        '<div class="cover-tipo-pill">{}</div>'
        '</div>'
        '<div class="cover-titulo">{}</div>'
        '<div class="cover-ubicacion">{}{}</div>'
        '{}'
        '</div>'
        '{}'
    ).format(
        hero_html, brandmark_html, photocount_html,
        tipo_op, precio_principal, precio_sec_html,
        ICO["home"],
        titulo,
        ICO["pin"], ubicacion,
        specs_html,
        '<div class="cover-desc-wrap"><div class="cover-desc-ttl">Descripción</div><div class="cover-desc">{}</div></div>'.format(desc) if desc else '<div style="flex:1"></div>'
    )

    # ── Páginas de galería (6 fotos por página, igual que en el frontend) ──
    gallery_fotos = foto_urls[1:]
    gallery_contents = []
    for i in range(0, len(gallery_fotos), 6):
        batch = gallery_fotos[i:i+6]
        batch = batch + [None] * (6 - len(batch))
        imgs = "".join(
            '<img src="{}" alt="foto"/>'.format(images_b64.get(u, u)) if u else '<div class="ph-empty"></div>'
            for u in batch
        )
        gallery_contents.append(
            '<div class="fp-kicker"><div class="fp-kicker-left"><div class="fp-kicker-ico">{}</div><h2>Galería fotográfica</h2></div>'
            '<div class="fp-kicker-id">{}</div></div>'
            '<div class="photo-grid">{}</div>'.format(ICO["photo"], ubicacion, imgs)
        )

    # ── Características agrupadas por categoría ──
    def char_item(icon, lbl, val):
        return '<div class="char-item"><div class="char-ico">{}</div><div class="char-txt"><div class="char-lbl">{}</div><div class="char-val">{}</div></div></div>'.format(icon, lbl, val)

    prec_rows = []
    prec_rows.append(char_item(ICO["swap"], "Operación", tipo_op))
    if precio_venta: prec_rows.append(char_item(ICO["tag"], "Precio de venta" if es_venta_renta else "Precio", precio_venta))
    if es_venta_renta and precio_renta: prec_rows.append(char_item(ICO["tag"], "Precio de renta", precio_renta + "/mes"))
    if not precio_venta and not es_venta_renta and precio_renta: pass  # ya cubierto como precio principal arriba

    dist_rows = []
    if tipo_inmueble: dist_rows.append(char_item(ICO["home"], "Tipo de inmueble", tipo_inmueble))
    if rec:  dist_rows.append(char_item(ICO["bed"], "Recámaras", rec))
    if ban:  dist_rows.append(char_item(ICO["bath"], "Baños completos", ban))
    if mban: dist_rows.append(char_item(ICO["toilet"], "Medios baños", mban))
    if niveles: dist_rows.append(char_item(ICO["levels"], "Niveles", niveles))
    if anio: dist_rows.append(char_item(ICO["calendar"], "Año de construcción", anio))

    sup_rows = []
    if fmt_m2(m2c): sup_rows.append(char_item(ICO["area"], "Superficie construida", fmt_m2(m2c)))
    if fmt_m2(m2t): sup_rows.append(char_item(ICO["land"], "Superficie de terreno", fmt_m2(m2t)))
    if parking: sup_rows.append(char_item(ICO["parking"], "Estacionamientos", parking))

    ub_rows = []
    if colonia: ub_rows.append(char_item(ICO["pin"], "Colonia", colonia))
    if ciudad:  ub_rows.append(char_item(ICO["pin"], "Ciudad", ciudad))
    if direccion: ub_rows.append(char_item(ICO["route"], "Dirección", direccion))
    if id_prop: ub_rows.append(char_item(ICO["tag"], "Clave", id_prop))

    def group_html(titulo_grupo, rows):
        if not rows:
            return ""
        return '<div class="chars-group"><div class="chars-group-ttl">{}</div><div class="chars-grid">{}</div></div>'.format(titulo_grupo, "".join(rows))

    amen_html = ""
    if amenids:
        items = "".join('<div class="amen-item">{}{}</div>'.format(ICO["sparkles"], a.get("name") or a) for a in amenids)
        amen_html = '<div class="chars-group amen-section"><div class="chars-group-ttl">Amenidades y extras</div><div class="amen-grid">{}</div></div>'.format(items)

    chars_content = (
        '<div class="fp-kicker"><div class="fp-kicker-left"><div class="fp-kicker-ico">{}</div><h2>Características del inmueble</h2></div>'
        '<div class="fp-kicker-id">{}</div></div>'
        '<div class="chars-body">{}{}{}{}{}</div>'
    ).format(
        ICO["list"], id_prop,
        group_html("Operación y precio", prec_rows),
        group_html("Distribución", dist_rows),
        group_html("Superficie y estacionamiento", sup_rows),
        group_html("Ubicación", ub_rows),
        amen_html,
    )

    all_contents = [cover_content] + gallery_contents + [chars_content]
    total_pages = len(all_contents)
    pages_html = "".join(
        '<div class="ficha-page">{}{}</div>'.format(content, footer(i + 1, total_pages))
        for i, content in enumerate(all_contents)
    )

    # ── Sistema de diseño ──
    # Los colores salen de brokr-theme.css vía theme_css_for_pdf(): este
    # archivo ya no los duplica. Cero JetBrains Mono, cero mayúsculas
    # decorativas.
    # Tokens desde brokr-theme.css. Radios y sombras propios del
    # documento: la ficha es un impreso, no una pantalla.
    CSS = theme_css_for_pdf(
        "--r:14px; --r-sm:8px; --r-lg:28px; --r-pill:999px;"
        "--shadow-sm:0 1px 3px rgba(0,20,59,.10),0 1px 2px rgba(0,20,59,.06);"
        "--shadow-lg:0 18px 44px rgba(0,20,59,.18),0 4px 12px rgba(0,20,59,.10);"
    ) + """
*{box-sizing:border-box;margin:0;padding:0;-webkit-print-color-adjust:exact!important;print-color-adjust:exact!important;color-adjust:exact!important}
html,body{width:210mm}
body{font-family:var(--font-sans);background:var(--paper);color:var(--ink);-webkit-font-smoothing:antialiased}
.ficha-page{position:relative;width:210mm;height:297mm;background:var(--paper);display:flex;flex-direction:column;overflow:hidden;page-break-after:always}
.ficha-page:last-child{page-break-after:avoid}

.fp-kicker{display:flex;align-items:center;justify-content:space-between;padding:14px 24px;border-bottom:1px solid var(--line)}
.fp-kicker-left{display:flex;align-items:center;gap:10px}
.fp-kicker-ico{width:20px;height:20px;color:var(--sky-blue);flex-shrink:0}.fp-kicker-ico svg{width:100%;height:100%}
.fp-kicker h2{font-family:var(--font-display);font-size:17px;font-weight:700;color:var(--ink);letter-spacing:-.02em}
.fp-kicker-id{font-size:11px;color:var(--mute-2)}

.cover-hero-wrap{width:100%;height:128mm;position:relative;flex-shrink:0;background:linear-gradient(135deg,var(--sky-navy),var(--ink-2))}
.cover-hero{width:100%;height:100%;object-fit:cover;display:block}
.cover-hero-placeholder{width:100%;height:100%;display:flex;align-items:center;justify-content:center}
.cover-hero-placeholder svg{width:56px;height:56px;color:rgba(255,255,255,.35)}
.cover-brandmark{position:absolute;top:16px;left:20px;height:20px}.cover-brandmark img{height:100%;width:auto;display:block}
.cover-photocount{position:absolute;top:16px;right:20px;background:rgba(5,32,60,.55);color:#fff;font-size:11px;font-weight:500;padding:5px 11px;border-radius:var(--r-pill);display:flex;align-items:center;gap:5px}
.cover-photocount svg{width:13px;height:13px}

.cover-card{margin:-22mm 16mm 0;background:var(--bone);border-radius:var(--r-lg);box-shadow:var(--shadow-lg);border:1px solid var(--line);padding:20px 24px 4px;position:relative;z-index:2}
.cover-card-top{display:flex;align-items:flex-start;justify-content:space-between;gap:16px;margin-bottom:14px}
.cover-badge{display:inline-flex;align-items:center;background:var(--sky-navy);color:#fff;font-size:12px;font-weight:600;padding:5px 12px;border-radius:var(--r-pill);margin-bottom:10px}
.cover-precio-block{display:flex;flex-direction:column}
.cover-precio{font-family:var(--font-display);font-size:34px;font-weight:700;letter-spacing:-.03em;color:var(--ink);line-height:1.05}
.cover-precio-sec{font-size:12.5px;color:var(--mute);margin-top:4px;font-weight:500}
.cover-precio-sec b{color:var(--ink-2);font-weight:600}
.cover-tipo-pill{flex-shrink:0;width:46px;height:46px;border-radius:var(--r);background:var(--paper-2);display:flex;align-items:center;justify-content:center;color:var(--sky-navy)}
.cover-tipo-pill svg{width:22px;height:22px}
.cover-titulo{font-family:var(--font-display);font-size:16px;font-weight:700;color:var(--ink);margin-bottom:5px;letter-spacing:-.015em}
.cover-ubicacion{font-size:12.5px;color:var(--mute);display:flex;align-items:center;gap:5px;padding-bottom:16px}
.cover-ubicacion svg{width:13px;height:13px;flex-shrink:0;color:var(--mute-2)}
.cover-specs{display:grid;grid-template-columns:repeat(var(--spec-cols,4),1fr);border-top:1px solid var(--line);margin:0 -24px;padding:0 24px}
.spec-item{padding:13px 6px 12px;text-align:center;border-right:1px solid var(--line)}
.spec-item:last-child{border-right:none}
.spec-ico{width:20px;height:20px;margin:0 auto 6px;color:var(--sky-blue)}.spec-ico svg{width:100%;height:100%}
.spec-val{font-family:var(--font-display);font-size:16px;font-weight:700;color:var(--ink);line-height:1.1;letter-spacing:-.02em}
.spec-lbl{font-size:10.5px;color:var(--mute);margin-top:3px;font-weight:500}
.cover-desc-wrap{padding:18px 24px 14px;flex:1}
.cover-desc-ttl{font-family:var(--font-display);font-size:13px;font-weight:700;color:var(--ink);margin-bottom:8px;letter-spacing:-.01em}
.cover-desc{font-size:11.5px;color:var(--ink-2);line-height:1.7}

.photo-grid{display:grid;grid-template-columns:1fr 1fr;grid-auto-rows:1fr;gap:4px;padding:4px;flex:1;overflow:hidden;background:var(--paper-2)}
.photo-grid img{width:100%;height:100%;object-fit:cover;display:block}
.photo-grid .ph-empty{width:100%;height:100%;background:var(--paper-2)}

.chars-body{padding:20px 24px 8px;flex:1}
.chars-group{margin-bottom:18px}
.chars-group-ttl{font-size:11px;font-weight:700;color:var(--mute);text-transform:uppercase;letter-spacing:.06em;margin-bottom:9px;padding-bottom:7px;border-bottom:1px solid var(--line)}
.chars-grid{display:grid;grid-template-columns:1fr 1fr;gap:10px}
.char-item{display:flex;align-items:center;gap:10px;padding:10px 12px;background:var(--paper-2);border-radius:var(--r-sm)}
.char-ico{width:18px;height:18px;color:var(--sky-blue);flex-shrink:0}.char-ico svg{width:100%;height:100%}
.char-txt{min-width:0}
.char-lbl{font-size:10px;color:var(--mute);margin-bottom:1px}
.char-val{font-size:13px;font-weight:600;color:var(--ink);letter-spacing:-.01em;overflow-wrap:anywhere}
.amen-grid{display:flex;flex-wrap:wrap;gap:7px}
.amen-item{display:inline-flex;align-items:center;gap:6px;font-size:11.5px;padding:6px 12px;background:var(--paper-2);border-radius:var(--r-pill);color:var(--ink-2);border:1px solid var(--line);font-weight:500}
.amen-item svg{width:12px;height:12px;color:var(--sky-blue);flex-shrink:0}

.ficha-footer{width:100%;height:42px;background:var(--sky-navy);display:flex;align-items:center;justify-content:space-between;padding:0 22px;flex-shrink:0;margin-top:auto}
.ft-brand{display:flex;align-items:center;gap:8px}
.ft-brand img{height:16px;width:auto;display:block;opacity:.95}
.ft-brand span{font-size:10px;font-weight:500;color:rgba(255,255,255,.6);letter-spacing:.01em}
.ft-meta{display:flex;align-items:center;gap:10px;font-size:10px;color:rgba(255,255,255,.5)}
.ft-id{letter-spacing:.03em}
@page{size:A4 portrait;margin:0}
"""

    return (
        "<!DOCTYPE html><html lang='es'><head><meta charset='UTF-8'/>"
        "<style>{}</style></head><body>{}</body></html>"
    ).format(CSS, pages_html)



# ────────────────────────────────────────────
# NOTICIAS INMOBILIARIAS — RSS REAL
# ────────────────────────────────────────────
import xml.etree.ElementTree as ET

@app.post("/ficha-pdf")
async def generar_ficha_pdf(p: dict, request: Request):
    """Generate PDF from property data dict using Playwright."""
    _uid = await get_user_id_from_token(request)
    exigir_cupo(request, _uid)
    exigir_sesion(request, _uid)
    import httpx
    
    # Collect all image URLs
    fotos = p.get("property_images") or []
    urls = list(set(filter(None, [f.get("url") or f.get("original") for f in fotos])))
    
    # Download all images concurrently and convert to base64
    images_b64 = {}
    async with httpx.AsyncClient(timeout=30) as client:
        async def fetch_img(url):
            try:
                r = await client.get(url, follow_redirects=True, timeout=10.0)
                if r.status_code == 200:
                    ext = url.split(".")[-1].split("?")[0].lower()
                    mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
                            "webp": "image/webp", "gif": "image/gif"}.get(ext, "image/jpeg")
                    b64 = base64.b64encode(r.content).decode()
                    images_b64[url] = f"data:{mime};base64,{b64}"
            except Exception:
                pass  # skip failed images, show blank

        # Limit to 19 gallery images (1 hero + 18 gallery = 3 full pages max)
        await asyncio.gather(*[fetch_img(u) for u in urls[:19]])
    
    # Build HTML
    html = build_ficha_html(p, images_b64)
    
    # Render to PDF with Playwright
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(args=["--no-sandbox", "--disable-dev-shm-usage"])
        page = await browser.new_page()
        # Use domcontentloaded instead of networkidle — images are already base64
        await page.set_content(html, wait_until="domcontentloaded")
        await page.wait_for_timeout(500)  # small wait for fonts
        pdf_bytes = await page.pdf(
            format="A4",
            print_background=True,
            margin={"top": "0", "right": "0", "bottom": "0", "left": "0"}
        )
        await browser.close()
    
    from fastapi.responses import JSONResponse
    import re as _re2
    id_prop   = p.get("public_id") or p.get("id") or ""
    loc       = p.get("location") or {}
    colonia   = (loc.get("name") or "").strip()
    tipo_raw  = (p.get("property_type") or "Propiedad").strip()
    # Sanitize: remove accents and special chars for filename
    def _slug(s):
        for a, b in [('á','a'),('é','e'),('í','i'),('ó','o'),('ú','u'),('ü','u'),('ñ','n'),
                     ('Á','A'),('É','E'),('Í','I'),('Ó','O'),('Ú','U'),('Ñ','N')]:
            s = s.replace(a, b)
        return _re2.sub(r'[^A-Za-z0-9_]', '_', s).strip('_')
    parts = ["Ficha"]
    if colonia:  parts.append(_slug(colonia))
    filename = "_".join(parts) + ".pdf"
    token = str(_uuid.uuid4()).replace("-","")[:16]
    _pdf_store[token] = (pdf_bytes, filename)
    # Clean old entries if too many
    if len(_pdf_store) > 50:
        oldest = list(_pdf_store.keys())[0]
        del _pdf_store[oldest]
    return JSONResponse({"token": token, "filename": filename})

# ════════════════════════════════════════════════════════════════
# META GRAPH API — capa común
# ════════════════════════════════════════════════════════════════
# Todas las llamadas al Graph API de Meta (Facebook) pasan por aquí.
# Antes cada endpoint hacía su propio httpx.get/post: la versión de la API
# estaba escrita a mano en ~40 lugares (y una se quedó en v18.0), nadie
# reintentaba cuando Meta contestaba 429, y los errores se devolvían como
# texto crudo. Esta capa arregla las tres cosas de un solo lugar.
#
# Es el espejo de _eb_get_reintentos() (EasyBroker), pero para Meta:
# Meta además codifica el motivo real del rechazo en `error.code`, no solo
# en el status HTTP, y publica su presupuesto de llamadas en la cabecera
# X-Business-Use-Case-Usage. Ambas cosas se honran abajo.

_fb_log = logging.getLogger("broquer.facebook")


# ─── Cifrado de tokens en reposo ──────────────────────────────────────────────
# Los tokens de Meta (página y usuario) vivían en texto plano en Supabase.
# Quien leyera esa tabla —un respaldo filtrado, una service_role key expuesta,
# un empleado con acceso a la consola— podía publicar y GASTAR en nombre del
# agente. Ahora se guardan cifrados con Fernet (AES-128-CBC + HMAC).
#
# Compatibilidad: los valores viejos siguen en claro y se leen igual. Se
# vuelven a escribir cifrados en cuanto la fila se actualiza, o de un jalón con
# POST /facebook/encrypt-tokens.
#
# Sin TOKEN_ENC_KEY configurada todo sigue funcionando en claro (y se avisa una
# vez en el log). Generar una llave:
#     python3 -c "from cryptography.fernet import Fernet; print(Fernet.generate_key().decode())"










# Códigos de error de Meta que significan "vuelve a intentar", NO "estás mal".
#   1     · API Unknown (error transitorio del lado de Meta)
#   2     · API Service (servicio temporalmente caído)
#   4     · Application request limit reached (límite de la app)
#   17    · User request limit reached (límite del usuario)
#   32    · Page-level throttling
#   341   · Application limit reached (límite temporal)
#   613   · Calls to this API have exceeded the rate limit
#   80000-80006 · Rate limits por caso de uso (80004 = ads_management)

# Códigos que significan "el token murió" — reintentar no sirve de nada.

# Interruptor de emergencia: si el appsecret_proof rompiera algo en producción
# se apaga con FB_APPSECRET_PROOF=0 en Railway sin tocar código.






# Errores de Meta traducidos a español de negocio. La llave es el
# error_subcode (o el code si no hay subcode) que manda Meta.
















# ─── Tokens: vida, permisos y avisos de expiración ────────────────────────────

# Los tokens de larga duración de Meta duran ~60 días. Cuando Meta no manda
# expires_in (tokens de página, que no expiran solos), asumimos este valor para
# poder avisar de todas formas.

# Días antes de la expiración en que empezamos a avisar en la UI.

# Permisos sin los cuales el módulo de anuncios no puede funcionar.






# ════════════════════════════════════════════════════════════════
# META — memoria de lo que Broquer creó (tabla fb_ad_entities)
# ════════════════════════════════════════════════════════════════
# Antes, crear un anuncio era una operación sin memoria: si el flujo se rompía
# a la mitad, los IDs se perdían y los recursos quedaban huérfanos en la cuenta
# publicitaria sin que nadie supiera que existían. Y un doble clic creaba dos
# campañas cobrando en paralelo.
#
# Todo esto degrada con elegancia: si la tabla no existe todavía (migración sin
# correr), se registra un aviso en el log y el anuncio se crea igual. Perder la
# bitácora no puede ser motivo para no poder anunciar.



def _sb_headers(extra: dict = None) -> dict:
    h = {"apikey": SUPABASE_SERVICE_KEY,
         "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
         "Content-Type": "application/json"}
    if extra:
        h.update(extra)
    return h












# ─── FACEBOOK OAUTH ───────────────────────────────────────────────────────────

# ────────────────────────────────────────────
# FACEBOOK — guardar / leer conexión por usuario
# ────────────────────────────────────────────





























# ─── FACEBOOK ADS ─────────────────────────────────────────────────────────────













# Periodos que Meta acepta en `date_preset`. Se valida contra esta lista para
# no reenviar a Meta cualquier cosa que llegue por query string.

# Breakdowns soportados. Meta no deja combinar cualquiera con cualquiera; esta
# lista es la que el módulo ofrece y sabe pintar.

# Las acciones que de verdad importan para un anuncio Click-to-Messenger.
# El KPI real del agente inmobiliario NO son las impresiones: son las
# conversaciones abiertas en Messenger y lo que cuesta cada una.









# Traducción de los effective_status de Meta. Un anuncio puede decir ACTIVE y
# no entregar nada porque Meta lo rechazó: sin esto el agente solo ve que "no
# llegan mensajes" y no sabe por qué.




# ════════════════════════════════════════════════════════════════
# META — Lead Ads: webhook y captura automática de prospectos
# ════════════════════════════════════════════════════════════════
# Un "Lead Ad" es el anuncio con formulario dentro de Facebook: la persona
# llena sus datos sin salir de la app. Meta avisa por webhook y hay que ir a
# recoger el lead con el token de la página.
#
# Sin esto, los leads se quedaban en Meta hasta que alguien se acordaba de
# bajarlos a mano — y un prospecto inmobiliario que espera dos días ya le
# compró a alguien más.

# Token que Meta usa para verificar la suscripción. Si no está configurado, el
# webhook queda cerrado (no se acepta ninguna suscripción a ciegas).
# Secreto para validar la firma. Se cae a FB_APP_SECRET porque los Lead Ads
# viven en la misma app de Meta que los anuncios.








# Cómo se llaman los campos estándar de Meta y a qué columna del CRM van.








# ════════════════════════════════════════════════════════════════
# META — públicos personalizados y similares (desde el CRM)
# ════════════════════════════════════════════════════════════════
# Sube los contactos del agente a Meta HASHEADOS (SHA-256) para poder
# anunciarle a su propia cartera, y para generar "públicos similares"
# (lookalikes): gente parecida a quienes ya le compraron.
#
# Meta NUNCA recibe datos en claro: el hash se hace aquí y es irreversible.
# Aun así, subir datos de clientes exige que el dueño de la cuenta haya
# aceptado las Condiciones de Públicos Personalizados en Business Manager;
# si no lo hizo, Meta rechaza con el código 2654 y aquí se traduce a
# instrucciones concretas en vez de un error críptico.

























# ════════════════════════════════════════════════════════════════
# META — AUTODIAGNÓSTICO (solo contra cuenta de PRUEBAS)
# ════════════════════════════════════════════════════════════════
# Ejercita la integración de punta a punta contra una TEST AD ACCOUNT de Meta:
# crea campaña, conjunto, creativo y anuncio de verdad, los lee, los prende y
# apaga, y al final los borra. Las cuentas de prueba de Meta NO cobran.
#
# Tres candados para que esto no pueda correr contra producción:
#   1. FB_QA_ENABLED=1 en el entorno.
#   2. FB_QA_AD_ACCOUNT_ID apuntando explícitamente a la cuenta de pruebas.
#   3. Verificación CONTRA META de que esa cuenta aparece en la lista de
#      cuentas de prueba de la app (/{app_id}/adaccounts). Si no aparece, se
#      aborta. No hay bandera para saltarse este candado.











# ════════════════════════════════════════════════════════════════
# STRIPE — SUSCRIPCIONES
# ════════════════════════════════════════════════════════════════


# IDs de Precios en Stripe (crear en dashboard.stripe.com → Productos → Precios)

# ── Broquer para Empresas ────────────────────────────────────────
# Se cobra en DOS líneas dentro de la misma suscripción de Stripe:
#   · base  → paquete de 5 usuarios, cantidad siempre 1
#   · extra → usuario adicional, cantidad = asientos - 5
# Así el dueño puede subir o bajar lugares sin cambiar de suscripción.


# Solo para pintar la pantalla. El cobro real siempre lo manda Stripe.




# ════════════════════════════════════════════════════════════════
# Contactos / Importar desde EasyBroker
# ════════════════════════════════════════════════════════════════


@app.post("/contactos/importar-archivo")
async def importar_contactos_archivo(request: Request, file: UploadFile = File(...)):
    """
    Importa contactos desde un archivo exportado de EasyBroker (o cualquier
    CSV / Excel con encabezados). Pensado para la migracion completa: la API
    de EasyBroker no expone toda la bitacora del CRM, pero el export de
    Contactos si trae notas, estatus, fechas reales y codigos de propiedad.

    Como funciona:
    - Acepta .csv (coma o punto y coma, UTF-8 o Latin-1) y .xlsx.
    - Detecta las columnas por nombre, sin importar el orden ni el idioma
      (Nombre, Telefono, Correo, Etiquetas, Fuente, Probabilidad, Notas,
      Estatus, Fecha de creacion, Agente, Propiedades, etc.).
    - Deduplica contra los contactos de la empresa por telefono y correo.
      En existentes solo rellena campos vacios; nunca pisa lo del usuario.
    - Conserva la fecha de creacion REAL del archivo para que Estadisticas
      muestre el historial en su mes correcto.
    - Detecta codigos EB-XXXX en cualquier columna o en las notas y liga el
      contacto con la propiedad ya importada via contactos_propiedades.
    """
    user_id = await get_user_id_from_token(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="Tu sesión expiró. Vuelve a iniciar sesión.")
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=500, detail="Supabase no está configurado en el servidor.")

    nombre_archivo = (file.filename or "").lower()
    contenido = await file.read()
    if not contenido:
        raise HTTPException(status_code=400, detail="El archivo llegó vacío.")
    if len(contenido) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="El archivo pesa más de 15 MB. Divide el export en partes más chicas.")

    # ─── Paso 1: leer filas del archivo como lista de dicts ───
    filas: list = []
    if nombre_archivo.endswith((".xlsx", ".xls")):
        try:
            import openpyxl
            from io import BytesIO
            wb = openpyxl.load_workbook(BytesIO(contenido), read_only=True, data_only=True)
            hoja = wb.worksheets[0]
            iterador = hoja.iter_rows(values_only=True)
            encabezados = None
            for row in iterador:
                celdas = ["" if v is None else str(v).strip() for v in row]
                if encabezados is None:
                    if not any(celdas):
                        continue
                    encabezados = celdas
                    continue
                if any(celdas):
                    filas.append(dict(zip(encabezados, celdas)))
            wb.close()
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"No se pudo leer el Excel: {str(e)[:150]}")
    else:
        import csv as _csv
        from io import StringIO
        texto = None
        for enc in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                texto = contenido.decode(enc)
                break
            except Exception:
                continue
        if texto is None:
            raise HTTPException(status_code=400, detail="No se pudo leer el archivo. Guárdalo como CSV UTF-8 o Excel.")
        primera = texto.splitlines()[0] if texto.splitlines() else ""
        delim = ";" if primera.count(";") > primera.count(",") else ","
        lector = _csv.DictReader(StringIO(texto), delimiter=delim)
        for row in lector:
            fila = {(k or "").strip(): ("" if v is None else str(v).strip()) for k, v in row.items()}
            if any(fila.values()):
                filas.append(fila)

    if not filas:
        raise HTTPException(status_code=400, detail="No se encontraron filas con datos. Revisa que la primera fila tenga los encabezados.")

    # ─── Paso 2: mapear encabezados por nombre, sin importar orden ni acentos ───
    import unicodedata

    def _norm(t: str) -> str:
        t = unicodedata.normalize("NFD", str(t or ""))
        t = "".join(c for c in t if unicodedata.category(c) != "Mn")
        return re.sub(r"[^a-z0-9 ]", "", t.lower()).strip()

    ALIAS = {
        "nombre":       ("nombre completo", "nombre", "name", "full name", "contacto", "cliente"),
        "apellido":     ("apellidos", "apellido", "last name"),
        "telefono":     ("telefono movil", "telefono celular", "telefonos", "telefono", "celular", "movil", "phone", "tel"),
        "wa":           ("whatsapp",),
        "email":        ("correo electronico", "correos", "correo", "email", "e mail", "mail"),
        "empresa":      ("empresa", "compania", "company"),
        "notas":        ("descripcion privada", "descripcion", "notas", "comentarios", "notes", "observaciones"),
        "etiquetas":    ("etiquetas", "tags"),
        "fuente":       ("fuente", "origen", "source"),
        "probabilidad": ("probabilidad", "probability"),
        "estatus":      ("estatus", "estado", "etapa", "status"),
        "calle":        ("direccion", "calle", "domicilio", "street"),
        "mpio":         ("municipio", "ciudad", "city"),
        "cp":           ("codigo postal", "cp", "postal code"),
        "fecha":        ("fecha de creacion", "fecha de registro", "fecha de alta", "creado", "created at", "fecha"),
        "agente":       ("agente asignado", "agente", "asesor", "responsable", "agent"),
        "props":        ("codigos de propiedad", "codigo de propiedad", "propiedades", "propiedades de interes", "propiedad", "inmuebles", "properties"),
        "tipo":         ("tipo de contacto", "tipo", "rol", "perfil"),
    }
    columnas_archivo = list(filas[0].keys())
    col_de = {}
    usadas = set()
    for campo, alias in ALIAS.items():
        for a in alias:
            for col in columnas_archivo:
                if col in usadas:
                    continue
                if a == _norm(col) or (len(a) > 3 and a in _norm(col)):
                    col_de[campo] = col
                    usadas.add(col)
                    break
            if campo in col_de:
                break

    if "nombre" not in col_de and "telefono" not in col_de and "email" not in col_de:
        raise HTTPException(status_code=400, detail=("No reconocí las columnas del archivo. Necesita al menos una de: "
                    "Nombre, Teléfono o Correo. Columnas recibidas: "
                    + ", ".join(columnas_archivo[:12])))

    _PROB = {"low": "baja", "baja": "baja", "medium": "media", "media": "media",
             "high": "alta", "alta": "alta"}
    _TIPO = {"comprador": "comprador", "buyer": "comprador",
             "vendedor": "vendedor", "seller": "vendedor",
             "propietario": "vendedor", "owner": "vendedor",
             "arrendador": "arrendador", "arrendatario": "arrendatario",
             "inquilino": "arrendatario"}
    _RE_EB = re.compile(r"EB-[A-Za-z0-9]{4,10}")

    def _tel_limpio_csv(x):
        t = re.sub(r"[^+\d]", "", str(x or ""))
        return t[:20] if len(t) >= 7 else ""

    def _fecha_iso(x):
        x = str(x or "").strip()
        if not x:
            return None
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d",
                    "%d/%m/%Y %H:%M", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(x[:19], fmt).isoformat()
            except Exception:
                continue
        return None

    def _valor(fila, campo):
        col = col_de.get(campo)
        return (fila.get(col) or "").strip() if col else ""

    # ─── Paso 3: universo existente de la empresa (dedupe org-wide) ───
    org_id_import = await get_org_id_for_user(user_id)
    filtro_org = ({"org_id": f"eq.{org_id_import}"} if org_id_import
                  else {"user_id": f"eq.{user_id}"})
    sb_headers = {
        "apikey": SUPABASE_SERVICE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        "Content-Type": "application/json",
    }
    prop_por_eb_id = {}
    pares_existentes = set()
    async with httpx.AsyncClient(timeout=20) as client:
        try:
            existentes = await get_rows(
                "contactos",
                {**filtro_org, "limit": "10000",
                 "select": "id,telefono,email,nombre,empresa,notas,fuente,probabilidad,calle,mpio,cp,wa,etiquetas,estatus"},
                timeout=20,
            )
        except httpx.HTTPStatusError:
            existentes = []
        try:
            propiedades_existentes = await get_rows(
                "propiedades",
                {**filtro_org, "eb_public_id": "not.is.null",
                 "select": "id,eb_public_id", "limit": "5000"},
                timeout=20,
            )
        except httpx.HTTPStatusError:
            propiedades_existentes = []
        for row in propiedades_existentes:
            if row.get("eb_public_id"):
                prop_por_eb_id[row["eb_public_id"]] = row["id"]
        try:
            vinculos_existentes = await get_rows(
                "contactos_propiedades",
                {"select": "contacto_id,propiedad_id", "limit": "20000"},
                timeout=20,
            )
        except httpx.HTTPStatusError:
            vinculos_existentes = []
        for v in vinculos_existentes:
            pares_existentes.add((v.get("contacto_id"), v.get("propiedad_id")))

    por_tel   = {_tel_limpio_csv(c.get("telefono")): c for c in existentes if _tel_limpio_csv(c.get("telefono"))}
    por_email = {(c.get("email") or "").strip().lower(): c for c in existentes if c.get("email")}

    mapa_ag = await _mapa_agentes_org(org_id_import, user_id)

    def _user_de_agente_txt(texto):
        """user_id de Broquer para el agente del archivo (correo o nombre)."""
        t = (texto or "").strip()
        if not t:
            return None
        if "@" in t:
            return mapa_ag["por_email"].get(t.lower())
        return mapa_ag["por_nombre"].get(mapa_ag["_nrm"](t))

    # ─── Paso 4: mapear, deduplicar y guardar ───
    importados = actualizados = omitidos = errores = 0
    vinculos_nuevos = 0
    sin_propiedad = 0

    async with httpx.AsyncClient(timeout=20) as client:
        for fila in filas:
            nombre = _valor(fila, "nombre")
            apellido = _valor(fila, "apellido")
            if apellido and apellido.lower() not in nombre.lower():
                nombre = f"{nombre} {apellido}".strip()
            nombre = nombre[:120]
            tel   = _tel_limpio_csv(_valor(fila, "telefono"))
            wa    = _tel_limpio_csv(_valor(fila, "wa"))
            email = _valor(fila, "email").lower()
            if email and ("@" not in email or " " in email):
                email = ""
            email = email[:120]
            if not nombre and not tel and not email:
                omitidos += 1
                continue

            notas = _valor(fila, "notas")[:2000]
            agente = _valor(fila, "agente")
            agente_uid = _user_de_agente_txt(agente)
            if agente and not agente_uid:
                # Sin match con un usuario de Broquer: al menos queda constancia
                linea = f"Asesor en EasyBroker: {agente}"
                notas = (notas + "\n" + linea).strip() if notas else linea
                notas = notas[:2000]
            etiquetas = [t.strip() for t in re.split(r"[,;|]", _valor(fila, "etiquetas")) if t.strip()][:40]
            fecha_real = _fecha_iso(_valor(fila, "fecha"))
            now_iso = datetime.utcnow().isoformat()

            m = {
                "nombre":       nombre,
                "telefono":     tel,
                "wa":           wa,
                "email":        email,
                "empresa":      _valor(fila, "empresa")[:120],
                "notas":        notas,
                "etiquetas":    etiquetas,
                "fuente":       (_valor(fila, "fuente") or "EasyBroker (archivo)")[:80],
                "probabilidad": _PROB.get(_valor(fila, "probabilidad").lower()),
                "estatus":      _valor(fila, "estatus").lower()[:40] or None,
                "calle":        _valor(fila, "calle")[:160],
                "mpio":         _valor(fila, "mpio")[:80],
                "cp":           _valor(fila, "cp")[:12],
            }

            # Codigos EB-XXXX: en la columna de propiedades y en las notas
            codigos = set(_RE_EB.findall(_valor(fila, "props")))
            codigos.update(_RE_EB.findall(notas))

            existente = (por_tel.get(tel) if tel else None) or (por_email.get(email) if email else None)

            if existente:
                contacto_id = existente["id"]
                patch = {}
                for campo in ("nombre", "telefono", "email", "wa", "empresa", "notas",
                              "fuente", "probabilidad", "estatus", "calle", "mpio", "cp"):
                    if not existente.get(campo) and m.get(campo):
                        patch[campo] = m[campo]
                if etiquetas:
                    prev = existente.get("etiquetas") or []
                    union = list(dict.fromkeys([*prev, *etiquetas]))
                    if union != prev:
                        patch["etiquetas"] = union
                if patch:
                    patch["updated_at"] = now_iso
                    try:
                        await patch_rows(
                            "contactos",
                            {"id": f"eq.{contacto_id}"},
                            patch,
                            timeout=20,
                            accepted_statuses=(200, 204),
                        )
                        actualizados += 1
                        existente.update(patch)
                    except httpx.HTTPStatusError:
                        errores += 1
                else:
                    omitidos += 1
            else:
                nuevo = {
                    "id":         str(_uuid.uuid4()),
                    "user_id":    agente_uid or user_id,
                    "org_id":     org_id_import,
                    "tipo":       _TIPO.get(_valor(fila, "tipo").lower(), "otro"),
                    "created_at": fecha_real or now_iso,
                    "updated_at": now_iso,
                    **m,
                }
                nuevo["nombre"] = nombre or "Sin nombre"
                nuevo = {k: v for k, v in nuevo.items() if v not in ("", None, [])}
                try:
                    await post_rows(
                        "contactos",
                        nuevo,
                        prefer="return=minimal",
                        timeout=20,
                        accepted_statuses=(200, 201, 204),
                    )
                    importados += 1
                    contacto_id = nuevo["id"]
                    if tel:
                        por_tel[tel] = {"id": contacto_id, **m}
                    if email:
                        por_email[email] = {"id": contacto_id, **m}
                except httpx.HTTPStatusError:
                    errores += 1
                    continue

            # Ligar propiedades por codigo EB (solo las ya importadas en Broquer)
            for cod in codigos:
                propiedad_id = prop_por_eb_id.get(cod)
                if not propiedad_id:
                    sin_propiedad += 1
                    continue
                if (contacto_id, propiedad_id) in pares_existentes:
                    continue
                try:
                    await post_rows(
                        "contactos_propiedades",
                        {"user_id": user_id, "contacto_id": contacto_id,
                         "propiedad_id": propiedad_id, "relacion": "interes"},
                        prefer="return=minimal",
                        timeout=20,
                        accepted_statuses=(200, 201, 204),
                    )
                    vinculos_nuevos += 1
                    pares_existentes.add((contacto_id, propiedad_id))
                except httpx.HTTPStatusError:
                    pass

    return {
        "ok": True,
        "filas":         len(filas),
        "importados":    importados,
        "actualizados":  actualizados,
        "omitidos":      omitidos,
        "vinculos":      vinculos_nuevos,
        "sin_propiedad": sin_propiedad,
        "errores":       errores,
        "columnas":      {k: v for k, v in col_de.items()},
    }


# ════════════════════════════════════════════════════════════════
# Migración completa EasyBroker como TRABAJO EN SEGUNDO PLANO
# El navegador ya no sostiene peticiones largas (se caían con cualquier
# corte o reinicio): inicia el trabajo y consulta el avance cada pocos
# segundos. El trabajo corre en el servidor y sobrevive a recargas de
# página. Los tres pasos se llaman internamente (localhost) reusando la
# lógica existente sin duplicarla.
# ════════════════════════════════════════════════════════════════

# ─────────────────────────────────────────────
# ADMIN
# Endpoints basados en rol (admin/equipo/agente) + activo (bool).
# El rol gobierna el acceso; las suscripciones de Stripe son solo para agentes.
# Solo accesibles si el caller tiene rol=admin (verificado vía service key).
# ─────────────────────────────────────────────



# ════════════════════════════════════════════════════════════════
# Eliminar cuenta y datos del usuario (acción irreversible)
# ════════════════════════════════════════════════════════════════
