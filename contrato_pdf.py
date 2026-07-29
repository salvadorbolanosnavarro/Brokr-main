# ──────────────────────────────────────────────────────────────────────────
# contrato_pdf.py · Broquer — De DOCX a PDF
# ──────────────────────────────────────────────────────────────────────────
# Convierte cualquier .docx a PDF leyéndolo con python-docx y volviéndolo a
# pintar como HTML, que luego Playwright imprime. Sin LibreOffice.
#
# POR QUÉ ASÍ Y NO REESCRIBIENDO LOS CONTRATOS
#   generar_contrato.py tiene el texto legal de la promesa y del arrendamiento
#   revisado y en producción. Reescribirlo como plantilla HTML significaría
#   mantener dos copias de las mismas cláusulas, y el día que se corrija una
#   se va a corregir en una sola. Así que no se toca: se genera el DOCX igual
#   que siempre y aquí se relee para imprimirlo.
#
#   El efecto secundario es el bueno: como el convertidor no sabe nada de
#   promesas ni arrendamientos, sirve igual para los machotes que sube el
#   propio agente. Un contrato suyo, con su membrete, también se puede mandar
#   a firmar.
#
# POR QUÉ NO LIBREOFFICE
#   Es la vía obvia y es la equivocada aquí: son ~500 MB en la imagen y varios
#   segundos de arranque en frío en cada conversión. Railway ya trae Chromium
#   instalado para el resto de los PDF de la plataforma; esto solo lo reutiliza.
#
# LO QUE ESTE CONVERTIDOR NO CONSERVA
#   Imágenes incrustadas, encabezados y pies de página de Word, y numeración
#   automática de listas. Para contratos de texto — que es todo lo que se firma
#   aquí — no aplica ninguno. Si algún día hay que respetar un membrete con
#   logo, este es el archivo que hay que ampliar.
# ──────────────────────────────────────────────────────────────────────────

import html
import logging
from typing import List, Dict, Any

log = logging.getLogger("broquer.contrato_pdf")

# Alineación de python-docx a CSS. El None de Word significa "hereda", y en
# un contrato lo que se hereda es justificado.
_ALINEACION = {0: "left", 1: "center", 2: "right", 3: "justify", None: "justify"}


def _runs_a_html(parrafo) -> str:
    """Un párrafo de Word es una lista de fragmentos con formato propio.
    Se respetan negritas, subrayado y cursiva, que es lo que usan los
    contratos para marcar las partes y los títulos de cláusula."""
    partes = []
    for run in parrafo.runs:
        txt = html.escape(run.text or "")
        if not txt:
            continue
        # Word guarda los saltos manuales dentro del propio run.
        txt = txt.replace("\n", "<br/>")
        if run.bold:
            txt = f"<strong>{txt}</strong>"
        if run.underline:
            txt = f"<u>{txt}</u>"
        if run.italic:
            txt = f"<em>{txt}</em>"
        partes.append(txt)
    return "".join(partes)


def _pt(valor, defecto: float = 0.0) -> float:
    try:
        return float(valor.pt) if valor is not None else defecto
    except Exception:
        return defecto


def _cm(valor, defecto: float = 0.0) -> float:
    try:
        return float(valor.cm) if valor is not None else defecto
    except Exception:
        return defecto


def docx_a_html(ruta_docx: str, titulo: str = "Documento") -> str:
    """Lee el .docx y devuelve el HTML listo para imprimir."""
    from docx import Document

    doc = Document(ruta_docx)
    bloques: List[str] = []

    for parrafo in doc.paragraphs:
        contenido = _runs_a_html(parrafo)
        pf = parrafo.paragraph_format

        if not contenido.strip():
            # Los párrafos vacíos de Word son espaciado deliberado del autor
            # del contrato. Se conservan, si no el documento se compacta y
            # deja de parecerse al que el agente ya conoce.
            bloques.append('<p class="vacio">&nbsp;</p>')
            continue

        estilos = [
            f"text-align:{_ALINEACION.get(parrafo.alignment, 'justify')}",
            f"margin-top:{_pt(pf.space_before):.1f}pt",
            f"margin-bottom:{_pt(pf.space_after, 6.0):.1f}pt",
        ]
        izq = _cm(pf.left_indent)
        if izq:
            estilos.append(f"margin-left:{izq:.2f}cm")
        primera = _cm(pf.first_line_indent)
        if primera:
            estilos.append(f"text-indent:{primera:.2f}cm")

        # El tamaño se toma del primer fragmento con tamaño explícito: en
        # estos contratos todos los runs de un párrafo comparten cuerpo.
        tam = None
        for run in parrafo.runs:
            if run.font is not None and run.font.size is not None:
                tam = _pt(run.font.size)
                break
        if tam:
            estilos.append(f"font-size:{tam:.1f}pt")

        bloques.append(f'<p style="{";".join(estilos)}">{contenido}</p>')

    # Las tablas son raras en estos contratos, pero los machotes del agente
    # sí las traen (tablas de pagos, inventarios de mobiliario).
    for tabla in doc.tables:
        filas = []
        for fila in tabla.rows:
            celdas = "".join(
                f"<td>{html.escape(c.text or '').replace(chr(10), '<br/>')}</td>"
                for c in fila.cells
            )
            filas.append(f"<tr>{celdas}</tr>")
        if filas:
            bloques.append(f'<table class="tbl">{"".join(filas)}</table>')

    return f"""<!DOCTYPE html><html lang="es"><head><meta charset="utf-8"/>
<title>{html.escape(titulo)}</title>
<style>
  @page {{ size: Letter; margin: 25mm 22mm; }}
  * {{ box-sizing: border-box; }}
  /* Mismas fuentes que genera_contrato.py pone en el .docx (Calibri y
     Cambria), para que el PDF que se firma se vea igual que el Word que el
     agente descarga. Carlito y Caladea son sus equivalentes libres en Linux
     y miden exactamente lo mismo, así que los saltos de página coinciden. */
  body {{
    font-family: Calibri, Carlito, Helvetica, Arial, sans-serif;
    font-size: 10.5pt; line-height: 1.5; color: #000; margin: 0;
    -webkit-font-smoothing: antialiased;
  }}
  p {{ margin: 0; orphans: 3; widows: 3; }}
  p.vacio {{ height: 6pt; }}
  u {{ text-underline-offset: 2px; }}
  .tbl {{ width: 100%; border-collapse: collapse; margin: 10pt 0; font-size: 10pt; }}
  .tbl td {{ border: 1px solid #000; padding: 5pt 7pt; vertical-align: top; }}
  strong {{ font-family: Cambria, Caladea, "Times New Roman", serif; }}
</style></head><body>
{''.join(bloques)}
</body></html>"""


async def docx_a_pdf(ruta_docx: str, titulo: str = "Documento") -> bytes:
    """El PDF impreso del .docx. Mismo Chromium que usan ISR, AVM y fichas."""
    from playwright.async_api import async_playwright

    contenido = docx_a_html(ruta_docx, titulo)
    async with async_playwright() as pw:
        navegador = await pw.chromium.launch(args=["--no-sandbox", "--disable-dev-shm-usage"])
        pagina = await navegador.new_page()
        await pagina.set_content(contenido, wait_until="domcontentloaded")
        await pagina.wait_for_timeout(300)
        pdf = await pagina.pdf(
            format="Letter",
            print_background=True,
            display_header_footer=True,
            header_template="<div></div>",
            footer_template=(
                '<div style="width:100%;font-size:8pt;color:#666;padding:0 22mm;'
                'font-family:Calibri,Carlito,Helvetica,Arial,sans-serif;text-align:center;">'
                '<span class="pageNumber"></span> de <span class="totalPages"></span>'
                "</div>"
            ),
            margin={"top": "22mm", "right": "22mm", "bottom": "18mm", "left": "22mm"},
        )
        await navegador.close()
    return pdf


# ── Quién firma cada contrato ─────────────────────────────────────────────
# El formulario de contratos ya tiene los nombres de las partes capturados.
# Volvérselos a pedir al agente para mandarlo a firmar sería pedirle que
# teclee dos veces lo mismo. Esto los saca del mismo diccionario de datos.
#
# El agente NO aparece en ninguno de los dos: en la promesa y en el
# arrendamiento es intermediario, no parte.
_PARTES: Dict[str, List[Dict[str, str]]] = {
    "promesa": [
        {"campo": "nombre_vendedor",   "rol": "promitente_vendedor"},
        {"campo": "nombre_comprador",  "rol": "promitente_comprador"},
    ],
    "arrendamiento": [
        {"campo": "nombre_arrendador",   "rol": "arrendador"},
        {"campo": "nombre_arrendatario", "rol": "arrendatario"},
        {"campo": "nombre_fiador",       "rol": "fiador"},
    ],
}


def partes_del_contrato(tipo: str, datos: Dict[str, Any]) -> List[Dict[str, str]]:
    """Los firmantes que se pueden deducir del contrato ya capturado."""
    salida = []
    for parte in _PARTES.get(tipo, []):
        nombre = str(datos.get(parte["campo"]) or "").strip()
        if not nombre:
            continue
        salida.append({"nombre": nombre, "rol": parte["rol"]})
    return salida
