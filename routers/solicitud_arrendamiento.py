from __future__ import annotations

from typing import List

from fastapi import APIRouter, File, Request, UploadFile


def create_router(get_context):
    router = APIRouter()

    @router.post("/solicitud-arrendamiento/analizar")
    async def analizar_solicitud_arrendamiento(
        request: Request,
        file: UploadFile = File(...),
        documentos: List[UploadFile] = File(default=[]),
    ):
        """
        Lee una solicitud de arrendamiento (PDF, imagen JPG/PNG/WEBP o DOCX) más
        hasta 5 documentos de respaldo opcionales (comprobantes de ingresos, escrituras
        del aval, INE, estados de cuenta, etc.) y los cruza todos con Claude Sonnet 4.6.
        Devuelve JSON estructurado con puntaje, riesgo, hallazgos y recomendaciones.
        Solicitud principal: máx 15 MB. Documentos adicionales: máx 8 MB c/u.
        Requiere usuario autenticado.
        """
        deps = get_context()
        get_user_id_from_token = deps["get_user_id_from_token"]
        HTTPException = deps["HTTPException"]
        ANTHROPIC_API_KEY = deps["ANTHROPIC_API_KEY"]
        ANTHROPIC_BASE = deps["ANTHROPIC_BASE"]
        _track_anthropic = deps["_track_anthropic"]
        httpx = deps["httpx"]
        base64 = deps["base64"]
        io = deps["io"]
        re = deps["re"]
        json = deps["json"]

        # Auth
        user_id = await get_user_id_from_token(request)
        if not user_id:
            raise HTTPException(status_code=401, detail="Inicia sesión para usar este módulo.")
        if not ANTHROPIC_API_KEY:
            raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY no configurada en el servidor.")

        # Leer archivo y validar tamaño
        content = await file.read()
        if len(content) > 15 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="Archivo demasiado grande (máx 15 MB).")
        if len(content) < 100:
            raise HTTPException(status_code=400, detail="Archivo vacío o corrupto.")

        fname = (file.filename or "").lower()
        ctype = (file.content_type or "").lower()

        is_pdf = ctype == "application/pdf" or fname.endswith(".pdf")
        is_docx = "wordprocessingml" in ctype or fname.endswith(".docx")
        is_image = (
            ctype.startswith("image/")
            or any(fname.endswith(x) for x in [".jpg", ".jpeg", ".png", ".webp", ".gif"])
        )

        # System prompt: rúbrica de evaluación + formato JSON estricto
        SYSTEM_PROMPT = """Eres un perito experto en evaluación de solicitudes de arrendamiento inmobiliario en México. Analizas con el rigor de un banco o inmobiliaria seria. Detectas inconsistencias, riesgos de impago y posibles fraudes.

Envuelve tu respuesta SIEMPRE entre las etiquetas <output> y </output>. Dentro de esas etiquetas coloca ÚNICAMENTE el JSON, sin texto adicional, sin bloques de markdown, sin comentarios. Así:
<output>
{ ... tu JSON aquí ... }
</output>

La estructura del JSON debe ser:
{
  "puntaje": <entero 0-100>,
  "nivel_riesgo": "verde" | "amarillo" | "rojo",
  "veredicto_corto": "<1-2 líneas resumiendo el caso>",
  "datos_extraidos": {
    "nombre_solicitante": "<string o null>",
    "edad": "<string o null>",
    "ocupacion": "<string o null>",
    "ingresos_mensuales_mxn": <número o null>,
    "renta_solicitada_mxn": <número o null>,
    "ratio_ingreso_renta": <número o null>,
    "tiene_aval": <true | false | null>,
    "tiene_referencias": <true | false | null>
  },
  "secciones": [
    {"categoria": "Identificación", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Domicilio", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Empleo e ingresos", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Estabilidad y referencias", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Fiador o garantía", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Indicadores PLD", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]},
    {"categoria": "Coherencia documental", "estatus": "ok"|"atencion"|"critico"|"faltante", "puntos": ["..."]}
  ],
  "banderas_rojas": ["..."],
  "recomendaciones": ["..."]
}

Rúbrica de puntaje:
- 90-100 (verde): completo, coherente, ratio ingreso/renta >= 3x, aval sólido con propiedad libre de gravamen
- 75-89 (verde): mayoritariamente completo, ratio 2.5-3x, mínimas faltantes
- 60-74 (amarillo): incompleto pero rescatable, ratio 2-2.5x o aval débil
- 40-59 (amarillo/rojo): faltan elementos críticos, ratio 1.5-2x, o referencias no verificables
- 0-39 (rojo): inconsistencias graves, posibles indicios de falsificación, datos críticos ausentes, ratio < 1.5x

Reglas estrictas:
1. Si no puedes extraer un dato, ponlo en null. NUNCA inventes información.
2. Calcula ratio_ingreso_renta = ingresos_mensuales_mxn / renta_solicitada_mxn cuando ambos estén presentes. Devuélvelo con 2 decimales.
3. En "secciones" SIEMPRE devuelve las 7 categorías en ese orden, aunque alguna esté "faltante".
4. estatus "faltante" = la solicitud simplemente no incluyó esa información (no es necesariamente malo, pero hay que pedirla).
5. estatus "critico" = riesgo grave detectado (no solo "falta", sino algo activamente alarmante).
6. Los "puntos" deben ser observaciones CONCRETAS, no generalidades. Cita datos específicos del documento cuando puedas.
7. "banderas_rojas" solo si hay riesgos genuinos: inconsistencias entre secciones, ratio < 2x sin aval, datos manipulados, referencias laborales sospechosas, fecha de emisión muy antigua, etc.
8. "recomendaciones" son acciones concretas que el agente debe hacer ANTES de firmar: verificar X comprobante con el patrón, confirmar Y referencia, pedir Z documento faltante, etc.
9. Indicadores PLD: revisa si hay coincidencias con criterios de actividad vulnerable de LFPIORPI (renta mensual >= 1,605 UMA = $188,282.55 MXN en 2026 obliga identificación del cliente; >= 3,210 UMA = $376,565 MXN obliga aviso al SAT)."""

        async def archivo_a_bloques(uf: UploadFile, etiqueta: str, max_bytes: int = 8 * 1024 * 1024):
            """Devuelve lista de bloques content para Claude según tipo de archivo."""
            raw = await uf.read()
            if len(raw) > max_bytes or len(raw) < 50:
                return []
            n = (uf.filename or "").lower()
            ct = (uf.content_type or "").lower()
            bloques = []
            bloques.append({"type": "text", "text": f"\n--- {etiqueta} ({uf.filename}) ---"})
            if ct == "application/pdf" or n.endswith(".pdf"):
                bloques.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": base64.standard_b64encode(raw).decode("utf-8")
                    }
                })
            elif "wordprocessingml" in ct or n.endswith(".docx"):
                try:
                    from docx import Document as _DocxDocument
                    _doc = _DocxDocument(io.BytesIO(raw))
                    _parts = []
                    for _p in _doc.paragraphs:
                        if _p.text and _p.text.strip():
                            _parts.append(_p.text.strip())
                    for _tbl in _doc.tables:
                        for _row in _tbl.rows:
                            for _cell in _row.cells:
                                for _p in _cell.paragraphs:
                                    if _p.text and _p.text.strip():
                                        _parts.append(_p.text.strip())
                    _txt = "\n".join(_parts)[:10000]
                    if _txt.strip():
                        bloques.append({"type": "text", "text": _txt})
                except Exception:
                    pass
            elif ct.startswith("image/") or any(n.endswith(x) for x in [".jpg", ".jpeg", ".png", ".webp"]):
                _mt = "image/jpeg"
                if n.endswith(".png") or "png" in ct:
                    _mt = "image/png"
                elif n.endswith(".webp") or "webp" in ct:
                    _mt = "image/webp"
                bloques.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": _mt,
                        "data": base64.standard_b64encode(raw).decode("utf-8")
                    }
                })
            return bloques

        user_content = []

        if is_pdf:
            b64 = base64.standard_b64encode(content).decode("utf-8")
            user_content.append({"type": "text", "text": "--- SOLICITUD DE ARRENDAMIENTO (documento principal) ---"})
            user_content.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": b64}
            })

        elif is_docx:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(content))
                parts = []
                for p in doc.paragraphs:
                    if p.text and p.text.strip():
                        parts.append(p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text and p.text.strip():
                                    parts.append(p.text.strip())
                extracted = "\n".join(parts)[:18000]
                if not extracted.strip():
                    raise HTTPException(status_code=400, detail="El DOCX no contiene texto legible.")
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"No se pudo leer el DOCX: {e}")
            user_content.append({
                "type": "text",
                "text": "--- SOLICITUD DE ARRENDAMIENTO (documento principal, formato Word) ---\n\n" + extracted
            })

        elif is_image:
            media_type = "image/jpeg"
            if fname.endswith(".png") or "png" in ctype:
                media_type = "image/png"
            elif fname.endswith(".webp") or "webp" in ctype:
                media_type = "image/webp"
            elif fname.endswith(".gif") or "gif" in ctype:
                media_type = "image/gif"
            b64 = base64.standard_b64encode(content).decode("utf-8")
            user_content.append({"type": "text", "text": "--- SOLICITUD DE ARRENDAMIENTO (documento principal) ---"})
            user_content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": media_type, "data": b64}
            })

        else:
            raise HTTPException(
                status_code=400,
                detail="Formato no soportado. Sube PDF, imagen (JPG/PNG/WEBP) o DOCX."
            )

        docs_validos = (documentos or [])[:5]
        nombres_extra = []
        for i, doc_extra in enumerate(docs_validos, start=1):
            etiqueta = f"DOCUMENTO DE RESPALDO #{i}"
            bloques = await archivo_a_bloques(doc_extra, etiqueta)
            if bloques:
                user_content.extend(bloques)
                nombres_extra.append(doc_extra.filename or f"documento_{i}")

        if nombres_extra:
            USER_INSTRUCTION = (
                f"Se adjuntan {len(nombres_extra)} documento(s) de respaldo además de la solicitud principal: "
                + ", ".join(nombres_extra) + ".\n"
                "Cruza la información de todos los documentos entre sí:\n"
                "- Verifica que los ingresos declarados en la solicitud coincidan con los comprobantes.\n"
                "- Verifica que el aval tenga solvencia real según su escritura u otro documento.\n"
                "- Detecta inconsistencias entre lo declarado en la solicitud y lo que muestran los respaldos.\n"
                "- Menciona discrepancias específicas en la sección 'Coherencia documental' y en banderas_rojas si aplica.\n\n"
                "Devuelve tu evaluación ÚNICAMENTE dentro de etiquetas <output></output>, "
                "como se indica en el system prompt. Solo JSON entre esas etiquetas."
            )
        else:
            USER_INSTRUCTION = (
                "Analiza esta solicitud de arrendamiento. "
                "Devuelve tu evaluación ÚNICAMENTE dentro de etiquetas <output></output>, "
                "como se indica en el system prompt. Solo JSON entre esas etiquetas, nada más."
            )

        user_content.append({"type": "text", "text": USER_INSTRUCTION})

        try:
            async with httpx.AsyncClient(timeout=150) as client:
                r = await client.post(
                    f"{ANTHROPIC_BASE}/messages",
                    headers={
                        "x-api-key": ANTHROPIC_API_KEY,
                        "anthropic-version": "2023-06-01",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": "claude-sonnet-4-6",
                        "max_tokens": 4096,
                        "system": SYSTEM_PROMPT,
                        "messages": [{"role": "user", "content": user_content}]
                    }
                )
            if r.status_code != 200:
                err_txt = (r.text or "")[:300]
                raise HTTPException(
                    status_code=502,
                    detail=f"Error Claude {r.status_code}: {err_txt}"
                )

            data = r.json()
            _track_anthropic(user_id, "solicitud-arr", "/solicitud-arrendamiento/analizar",
                             data, modelo=data.get("model") or "claude-sonnet-4-6")
            reply_text = ""
            try:
                reply_text = data.get("content", [{}])[0].get("text", "")
            except Exception:
                pass
            if not reply_text:
                raise HTTPException(status_code=502, detail="Claude devolvió respuesta vacía.")

            json_str = None
            tag_match = re.search(r'<output>\s*(.*?)\s*</output>', reply_text, re.DOTALL | re.IGNORECASE)
            if tag_match:
                json_str = tag_match.group(1).strip()
            else:
                brace_match = re.search(r'\{.*\}', reply_text, re.DOTALL)
                if brace_match:
                    json_str = brace_match.group().strip()

            if not json_str:
                raise HTTPException(status_code=502, detail="Claude no devolvió JSON válido.")

            json_str = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', json_str)
            json_str = json_str.lstrip('\ufeff')

            try:
                parsed = json.loads(json_str)
            except json.JSONDecodeError as e:
                try:
                    json_str2 = re.sub(
                        r'(?<=[:{,\[])\s*"((?:[^"\\]|\\.)*)"\s*(?=[,}\]:])',
                        lambda m: '"' + m.group(1).replace('"', '\\"') + '"',
                        json_str
                    )
                    parsed = json.loads(json_str2)
                except Exception:
                    raise HTTPException(
                        status_code=502,
                        detail=f"JSON inválido de Claude: {str(e)[:120]}"
                    )

            if "puntaje" not in parsed or "nivel_riesgo" not in parsed:
                raise HTTPException(status_code=502, detail="Respuesta sin estructura esperada.")

            parsed.setdefault("datos_extraidos", {})
            parsed.setdefault("secciones", [])
            parsed.setdefault("banderas_rojas", [])
            parsed.setdefault("recomendaciones", [])
            parsed.setdefault("veredicto_corto", "")

            return parsed

        except HTTPException:
            raise
        except httpx.TimeoutException:
            raise HTTPException(status_code=504, detail="El análisis tardó demasiado. Intenta de nuevo.")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error procesando: {str(e)[:200]}")

    return router
